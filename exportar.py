import io
import os
import re
import pandas as pd
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from docx import Document
from docx.shared import Pt, Cm as DCm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Registro de fuentes Calibri para PDF (opcional, fallback a Helvetica) ─────
_PDF_FONT      = "Helvetica"
_PDF_FONT_BOLD = "Helvetica-Bold"
try:
    from reportlab.pdfbase import pdfmetrics as _pdfm
    from reportlab.pdfbase.ttfonts import TTFont as _TTFont
    _cal_r = next((p for p in [
        r"C:\Windows\Fonts\calibri.ttf",
        "/usr/share/fonts/truetype/msttcorefonts/Calibri.ttf",
        "/usr/share/fonts/calibri/Calibri.ttf",
    ] if os.path.exists(p)), None)
    _cal_b = next((p for p in [
        r"C:\Windows\Fonts\calibrib.ttf",
        "/usr/share/fonts/truetype/msttcorefonts/Calibri_Bold.ttf",
        "/usr/share/fonts/calibri/Calibrib.ttf",
    ] if os.path.exists(p)), None)
    if _cal_r:
        _pdfm.registerFont(_TTFont("Calibri", _cal_r))
        _PDF_FONT = "Calibri"
    if _cal_b:
        _pdfm.registerFont(_TTFont("Calibri-Bold", _cal_b))
        _PDF_FONT_BOLD = "Calibri-Bold"
except Exception:
    pass


def _fmt(n: float) -> str:
    """Formato argentino: 1.234.567,89"""
    return f"{n:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")


def exportar_excel(df: pd.DataFrame) -> bytes:
    buf = io.BytesIO()
    out = pd.DataFrame({
        "Período": df["periodo"],
        "Capital ($)": df["capital"],
        "Intereses desde": df["fecha_desde"].dt.strftime("%d/%m/%Y"),
        "Fecha pago": df["fecha_pago"].dt.strftime("%d/%m/%Y"),
        "Índice inicial": df["indice_inicial"],
        "Índice final": df["indice_final"],
        "Coeficiente": df["coeficiente"],
        "Interés moratorio ($)": df["interes"],
        "Total ($)": df["total"],
    })

    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        out.to_excel(writer, index=False, sheet_name="Liquidación")
        ws = writer.sheets["Liquidación"]
        for col in ws.columns:
            max_len = max(len(str(cell.value or "")) for cell in col)
            ws.column_dimensions[col[0].column_letter].width = min(max_len + 3, 30)

    return buf.getvalue()


def exportar_pdf(df: pd.DataFrame, titulo: str = "PLANILLA DE LIQUIDACIÓN - INTERESES MORATORIOS - DOCTRINA RASTRILLA") -> bytes:
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=landscape(A4),
        leftMargin=1 * cm, rightMargin=1 * cm,
        topMargin=1.5 * cm, bottomMargin=1.5 * cm,
    )
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph(titulo, styles["Title"]))
    story.append(Paragraph(
        "Tasa pasiva BCRA Com. 14290 · Cálculo por coeficiente acumulado · Criterio RASTRILLA",
        styles["Normal"],
    ))
    story.append(Spacer(1, 0.5 * cm))

    headers = ["Mes/Año", "Int. desde", "Dif. neta", "Índ. inicial", "Índ. final", "Coeficiente", "Interés moratorio", "Total"]
    rows = [headers]

    for _, row in df.iterrows():
        rows.append([
            row["periodo"],
            row["fecha_desde"].strftime("%d/%m/%Y"),
            f"$ {_fmt(row['capital'])}",
            f"{row['indice_inicial']:,.4f}",
            f"{row['indice_final']:,.4f}",
            f"{row['coeficiente']:.6f}",
            f"$ {_fmt(row['interes'])}",
            f"$ {_fmt(row['total'])}",
        ])

    total_capital = df["capital"].sum()
    total_interes = df["interes"].sum()
    total_general = df["total"].sum()
    rows.append(["TOTAL GENERAL", "", f"$ {_fmt(total_capital)}", "", "", "", f"$ {_fmt(total_interes)}", f"$ {_fmt(total_general)}"])

    available_width = 27.7 * cm
    col_widths = [w * cm for w in [2.8, 2.5, 3.8, 3.2, 3.2, 3.2, 4.0, 4.0]]
    t = Table(rows, colWidths=col_widths, repeatRows=1)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1e3a5f")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, 0), 8),
        ("ALIGN", (0, 0), (-1, 0), "CENTER"),
        ("ALIGN", (2, 1), (2, -1), "RIGHT"),
        ("ALIGN", (6, 1), (7, -1), "RIGHT"),
        ("FONTSIZE", (0, 1), (-1, -1), 8),
        ("ROWBACKGROUNDS", (0, 1), (-1, -2), [colors.white, colors.HexColor("#eef2f7")]),
        ("BACKGROUND", (0, -1), (-1, -1), colors.HexColor("#fef3cd")),
        ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("ALIGN", (0, -1), (-1, -1), "CENTER"),
    ]))
    story.append(t)

    doc.build(story)
    return buf.getvalue()


# ── Ejecución de Sentencia — Excel (2 hojas) ───────────────────────────────────

def exportar_excel_ejecucion(resultado: dict) -> bytes:
    """Excel con dos hojas: Tramo A (desglose + cálculo único) y Tramo B."""
    buf = io.BytesIO()

    dia_120          = resultado["dia_120"]
    dia_121          = resultado["dia_121"]
    filas_a          = resultado["filas_a"]
    capital_a_total  = resultado["capital_a_total"]
    res_a            = resultado["resultado_a"]
    res_b            = resultado["resultado_b"]
    fecha_hasta      = resultado["fecha_hasta"]

    with pd.ExcelWriter(buf, engine="openpyxl") as writer:

        # ── Hoja Tramo A ──────────────────────────────────────────────────────
        # Tabla 1: aporte proporcional por período
        df_periodos = pd.DataFrame(
            [{"Período": f["periodo"], "Capital proporcional ($)": f["capital"]} for f in filas_a]
            + [{"Período": "TOTAL TRAMO A", "Capital proporcional ($)": capital_a_total}]
        )
        df_periodos.to_excel(writer, index=False, sheet_name="Tramo A", startrow=0)

        # Tabla 2: cálculo único de intereses sobre el total
        startrow_calc = len(df_periodos) + 3
        if not res_a.get("error"):
            df_calc = pd.DataFrame([{
                "Capital total ($)":     capital_a_total,
                "Intereses desde":       dia_121.strftime("%d/%m/%Y"),
                "Intereses hasta":       fecha_hasta.strftime("%d/%m/%Y"),
                "T₀ (Día 120)":          dia_120.strftime("%d/%m/%Y"),
                "Índice T₀":             res_a["indice_inicial"],
                "Índice Tₘ":             res_a["indice_final"],
                "Coeficiente":           res_a["coeficiente"],
                "Interés moratorio ($)": res_a["interes"],
                "Total ($)":             res_a["total"],
            }])
        else:
            df_calc = pd.DataFrame([{"Error": res_a["error"]}])
        df_calc.to_excel(writer, index=False, sheet_name="Tramo A", startrow=startrow_calc)

        # Tabla 3: resultado final (intereses A + B) y fecha hasta
        int_a_xls = float(res_a.get("interes", 0) or 0) if not res_a.get("error") else 0.0
        int_b_xls = float(res_b["interes"].sum()) if not res_b.empty and "interes" in res_b.columns else 0.0
        startrow_final = startrow_calc + len(df_calc) + 3
        df_final = pd.DataFrame([
            {
                "Concepto": "RESULTADO FINAL — Suma intereses moratorios Tramo A + Tramo B",
                "Valor":    int_a_xls + int_b_xls,
            },
            {
                "Concepto": "Calculado hasta — Efectivo pago conforme recibo que consta en autos",
                "Valor":    fecha_hasta.strftime("%d/%m/%Y"),
            },
        ])
        df_final.to_excel(writer, index=False, sheet_name="Tramo A", startrow=startrow_final)

        ws_a = writer.sheets["Tramo A"]
        for col in ws_a.columns:
            max_len = max(len(str(cell.value or "")) for cell in col)
            ws_a.column_dimensions[col[0].column_letter].width = min(max_len + 3, 55)

        # ── Hoja Tramo B ──────────────────────────────────────────────────────
        if not res_b.empty:
            df_b_ok = res_b[res_b["error"].isna()].copy()
            df_out = pd.DataFrame({
                "Período":               df_b_ok["periodo"],
                "Capital ($)":           df_b_ok["capital"],
                "Intereses desde":       df_b_ok["fecha_desde"].dt.strftime("%d/%m/%Y"),
                "Fecha pago":            df_b_ok["fecha_pago"].dt.strftime("%d/%m/%Y"),
                "Índice T₀":             df_b_ok["indice_inicial"],
                "Índice Tₘ":             df_b_ok["indice_final"],
                "Coeficiente":           df_b_ok["coeficiente"],
                "Interés moratorio ($)": df_b_ok["interes"],
                "Total ($)":             df_b_ok["total"],
            })
        else:
            df_out = pd.DataFrame(columns=[
                "Período", "Capital ($)", "Intereses desde", "Fecha pago",
                "Índice T₀", "Índice Tₘ", "Coeficiente", "Interés moratorio ($)", "Total ($)",
            ])

        df_out.to_excel(writer, index=False, sheet_name="Tramo B")
        ws_b = writer.sheets["Tramo B"]
        for col in ws_b.columns:
            max_len = max(len(str(cell.value or "")) for cell in col)
            ws_b.column_dimensions[col[0].column_letter].width = min(max_len + 3, 30)

    return buf.getvalue()


# ── Ejecución de Sentencia — PDF ───────────────────────────────────────────────

def exportar_pdf_ejecucion(
    resultado: dict,
    titulo: str = "EJECUCIÓN DE SENTENCIA — DOCTRINA RASTRILLA · VEGA",
) -> bytes:
    """PDF con sección Tramo A (desglose + cálculo único) y sección Tramo B."""
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=landscape(A4),
        leftMargin=1 * cm, rightMargin=1 * cm,
        topMargin=1.5 * cm, bottomMargin=1.5 * cm,
    )
    styles = getSampleStyleSheet()
    story  = []

    dia_120         = resultado["dia_120"]
    dia_121         = resultado["dia_121"]
    filas_a         = resultado["filas_a"]
    capital_a_total = resultado["capital_a_total"]
    res_a           = resultado["resultado_a"]
    res_b           = resultado["resultado_b"]
    fecha_hasta     = resultado["fecha_hasta"]

    # ── Encabezado ─────────────────────────────────────────────────────────────
    story.append(Paragraph(titulo, styles["Title"]))
    story.append(Paragraph(
        "Tasa pasiva BCRA Com. 14290 · Cálculo por coeficiente acumulado · Criterio RASTRILLA",
        styles["Normal"],
    ))
    story.append(Spacer(1, 0.4 * cm))

    # ── Tramo A ─────────────────────────────────────────────────────────────────
    story.append(Paragraph(
        f"<b>TRAMO A</b> — Períodos dentro del plazo de 120 días hábiles judiciales &nbsp;|&nbsp; "
        f"Día 120: <b>{dia_120.strftime('%d/%m/%Y')}</b> &nbsp;·&nbsp; "
        f"Intereses desde (Día 121): <b>{dia_121.strftime('%d/%m/%Y')}</b> &nbsp;·&nbsp; "
        f"Hasta: <b>{fecha_hasta.strftime('%d/%m/%Y')}</b>",
        styles["Normal"],
    ))
    story.append(Spacer(1, 0.2 * cm))

    # Tabla de aportes por período
    if filas_a:
        rows_a = [["Período", "Capital proporcional ($)"]]
        for f in filas_a:
            rows_a.append([f["periodo"], f"$ {_fmt(f['capital'])}"])
        rows_a.append(["TOTAL TRAMO A", f"$ {_fmt(capital_a_total)}"])
        t_a1 = Table(rows_a, colWidths=[4 * cm, 5 * cm])
        t_a1.setStyle(TableStyle([
            ("BACKGROUND",    (0, 0), (-1, 0),  colors.HexColor("#1e3a5f")),
            ("TEXTCOLOR",     (0, 0), (-1, 0),  colors.white),
            ("FONTNAME",      (0, 0), (-1, 0),  "Helvetica-Bold"),
            ("FONTSIZE",      (0, 0), (-1, -1), 8),
            ("ALIGN",         (1, 0), (1, -1),  "RIGHT"),
            ("ROWBACKGROUNDS",(0, 1), (-1, -2), [colors.white, colors.HexColor("#eef2f7")]),
            ("BACKGROUND",    (0, -1), (-1, -1), colors.HexColor("#fef3cd")),
            ("FONTNAME",      (0, -1), (-1, -1), "Helvetica-Bold"),
            ("GRID",          (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
            ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING",    (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(t_a1)
        story.append(Spacer(1, 0.2 * cm))

    # Tabla del cálculo único de Tramo A
    if not res_a.get("error"):
        rows_calc = [
            ["Capital total ($)", "Int. desde", "Int. hasta",
             "Índ. T₀", "Índ. Tₘ", "Coeficiente", "Interés moratorio ($)", "Total ($)"],
            [
                f"$ {_fmt(capital_a_total)}",
                dia_121.strftime("%d/%m/%Y"),
                fecha_hasta.strftime("%d/%m/%Y"),
                f"{res_a['indice_inicial']:,.4f}",
                f"{res_a['indice_final']:,.4f}",
                f"{res_a['coeficiente']:.6f}",
                f"$ {_fmt(res_a['interes'])}",
                f"$ {_fmt(res_a['total'])}",
            ],
        ]
        col_w_calc = [w * cm for w in [3.5, 2.5, 2.5, 3.0, 3.0, 3.0, 4.5, 4.5]]
        t_a2 = Table(rows_calc, colWidths=col_w_calc)
        t_a2.setStyle(TableStyle([
            ("BACKGROUND",    (0, 0), (-1, 0),  colors.HexColor("#2d6a4f")),
            ("TEXTCOLOR",     (0, 0), (-1, 0),  colors.white),
            ("FONTNAME",      (0, 0), (-1, 0),  "Helvetica-Bold"),
            ("FONTSIZE",      (0, 0), (-1, -1), 8),
            ("ALIGN",         (0, 0), (-1, -1), "CENTER"),
            ("BACKGROUND",    (0, 1), (-1, 1),  colors.HexColor("#d8f3dc")),
            ("FONTNAME",      (0, 1), (-1, 1),  "Helvetica-Bold"),
            ("GRID",          (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
            ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING",    (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(t_a2)
    else:
        story.append(Paragraph(f"Sin datos para Tramo A: {res_a.get('error', '')}", styles["Normal"]))

    story.append(Spacer(1, 0.5 * cm))

    # ── Tramo B ─────────────────────────────────────────────────────────────────
    story.append(Paragraph(
        "<b>TRAMO B</b> — Períodos posteriores al día 120 · Intereses por período (criterio Ampliación)",
        styles["Normal"],
    ))
    story.append(Spacer(1, 0.2 * cm))

    if not res_b.empty:
        df_b_ok = res_b[res_b["error"].isna()].copy()
        rows_b = [["Período", "Int. desde", "Capital ($)",
                   "Índ. T₀", "Índ. Tₘ", "Coeficiente", "Interés ($)", "Total ($)"]]
        for _, row in df_b_ok.iterrows():
            rows_b.append([
                row["periodo"],
                row["fecha_desde"].strftime("%d/%m/%Y"),
                f"$ {_fmt(row['capital'])}",
                f"{row['indice_inicial']:,.4f}",
                f"{row['indice_final']:,.4f}",
                f"{row['coeficiente']:.6f}",
                f"$ {_fmt(row['interes'])}",
                f"$ {_fmt(row['total'])}",
            ])
        t_cap_b = df_b_ok["capital"].sum()
        t_int_b = df_b_ok["interes"].sum()
        t_tot_b = df_b_ok["total"].sum()
        rows_b.append(["TOTAL TRAMO B", "", f"$ {_fmt(t_cap_b)}", "", "",
                        "", f"$ {_fmt(t_int_b)}", f"$ {_fmt(t_tot_b)}"])

        col_w_b = [w * cm for w in [2.8, 2.5, 3.8, 3.2, 3.2, 3.2, 4.0, 4.0]]
        t_b = Table(rows_b, colWidths=col_w_b, repeatRows=1)
        t_b.setStyle(TableStyle([
            ("BACKGROUND",    (0, 0),  (-1, 0),  colors.HexColor("#1e3a5f")),
            ("TEXTCOLOR",     (0, 0),  (-1, 0),  colors.white),
            ("FONTNAME",      (0, 0),  (-1, 0),  "Helvetica-Bold"),
            ("FONTSIZE",      (0, 0),  (-1, -1), 8),
            ("ALIGN",         (0, 0),  (-1, 0),  "CENTER"),
            ("ALIGN",         (2, 1),  (2, -1),  "RIGHT"),
            ("ALIGN",         (6, 1),  (7, -1),  "RIGHT"),
            ("ROWBACKGROUNDS",(0, 1),  (-1, -2), [colors.white, colors.HexColor("#eef2f7")]),
            ("BACKGROUND",    (0, -1), (-1, -1), colors.HexColor("#fef3cd")),
            ("FONTNAME",      (0, -1), (-1, -1), "Helvetica-Bold"),
            ("GRID",          (0, 0),  (-1, -1), 0.5, colors.HexColor("#cccccc")),
            ("VALIGN",        (0, 0),  (-1, -1), "MIDDLE"),
            ("TOPPADDING",    (0, 0),  (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0),  (-1, -1), 4),
            ("ALIGN",         (0, -1), (-1, -1), "CENTER"),
        ]))
        story.append(t_b)
    else:
        story.append(Paragraph("No hay períodos en Tramo B.", styles["Normal"]))

    # ── Resultado final ──────────────────────────────────────────────────────────
    story.append(Spacer(1, 0.6 * cm))
    int_a_val = float(resultado.get("resultado_a", {}).get("interes", 0) or 0)
    int_b_val = float(res_b["interes"].sum()) if not res_b.empty and "interes" in res_b.columns else 0.0
    int_total = int_a_val + int_b_val

    rows_final = [
        [
            "RESULTADO FINAL — Suma intereses moratorios Tramo A + Tramo B",
            f"$ {_fmt(int_total)}",
        ],
        [
            "Calculado hasta — Efectivo pago conforme recibo que consta en autos",
            fecha_hasta.strftime("%d/%m/%Y"),
        ],
    ]
    available_w = 27.7 * cm
    t_final = Table(rows_final, colWidths=[available_w * 0.72, available_w * 0.28])
    t_final.setStyle(TableStyle([
        ("BACKGROUND",    (0, 0), (-1, 0),  colors.HexColor("#052e16")),
        ("TEXTCOLOR",     (0, 0), (-1, 0),  colors.white),
        ("FONTNAME",      (0, 0), (-1, 0),  "Helvetica-Bold"),
        ("FONTSIZE",      (0, 0), (-1, 0),  10),
        ("BACKGROUND",    (0, 1), (-1, 1),  colors.HexColor("#f0fdf4")),
        ("TEXTCOLOR",     (0, 1), (-1, 1),  colors.HexColor("#052e16")),
        ("FONTNAME",      (0, 1), (-1, 1),  "Helvetica-Bold"),
        ("FONTSIZE",      (0, 1), (-1, 1),  9),
        ("ALIGN",         (1, 0), (1, -1),  "RIGHT"),
        ("GRID",          (0, 0), (-1, -1), 0.5, colors.HexColor("#16a34a")),
        ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING",    (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("LEFTPADDING",   (0, 0), (-1, -1), 8),
    ]))
    story.append(t_final)

    doc.build(story)
    return buf.getvalue()


# ── Intereses hasta el Cobro — DOCX (Escrito judicial) ───────────────────────

def _numero_a_palabras(monto: float) -> str:
    """Convierte monto a palabras en castellano (mayúsculas) para escritos judiciales.
    Ej: 250000.50 → 'DOSCIENTOS CINCUENTA MIL CON 50/100'
    """
    _UNI = ["", "UN", "DOS", "TRES", "CUATRO", "CINCO",
            "SEIS", "SIETE", "OCHO", "NUEVE"]
    _ESP = ["DIEZ", "ONCE", "DOCE", "TRECE", "CATORCE", "QUINCE",
            "DIECISÉIS", "DIECISIETE", "DIECIOCHO", "DIECINUEVE"]
    _DEC = ["", "DIEZ", "VEINTE", "TREINTA", "CUARENTA", "CINCUENTA",
            "SESENTA", "SETENTA", "OCHENTA", "NOVENTA"]
    _CEN = ["", "CIENTO", "DOSCIENTOS", "TRESCIENTOS", "CUATROCIENTOS",
            "QUINIENTOS", "SEISCIENTOS", "SETECIENTOS", "OCHOCIENTOS", "NOVECIENTOS"]
    _VTI = {21: "VEINTIÚN", 22: "VEINTIDÓS", 23: "VEINTITRÉS", 24: "VEINTICUATRO",
            25: "VEINTICINCO", 26: "VEINTISÉIS", 27: "VEINTISIETE",
            28: "VEINTIOCHO", 29: "VEINTINUEVE"}

    def _bloque(n: int) -> str:
        if n == 0:
            return ""
        if n == 100:
            return "CIEN"
        partes = []
        c, resto = divmod(n, 100)
        if c:
            partes.append(_CEN[c])
        if 10 <= resto <= 19:
            partes.append(_ESP[resto - 10])
        elif 21 <= resto <= 29:
            partes.append(_VTI[resto])
        elif resto == 20:
            partes.append("VEINTE")
        else:
            d, u = divmod(resto, 10)
            if d:
                partes.append(_DEC[d])
            if u:
                if d:
                    partes.append("Y")
                partes.append(_UNI[u])
        return " ".join(p for p in partes if p)

    entero = int(monto)
    cents  = round((monto - entero) * 100)

    if entero == 0:
        palabras = "CERO"
    else:
        millones   = entero // 1_000_000
        resto_mill = entero % 1_000_000
        miles      = resto_mill // 1_000
        resto_mil  = resto_mill % 1_000
        partes = []
        if millones == 1:
            partes.append("UN MILLÓN")
        elif millones > 1:
            partes.append(f"{_bloque(millones)} MILLONES")
        if miles == 1:
            partes.append("MIL")
        elif miles > 1:
            partes.append(f"{_bloque(miles)} MIL")
        if resto_mil:
            partes.append(_bloque(resto_mil))
        palabras = " ".join(p for p in partes if p)

    if cents:
        palabras += f" CON {cents:02d}/100"
    return palabras


def _limpiar_caratula(caratula: str) -> str:
    """Elimina '(agregar/quitar...)' y texto similar generado por el sistema judicial."""
    return re.sub(r"\s*\(agregar/quitar[^)]*\)", "", caratula, flags=re.IGNORECASE).strip()


def limpiar_expediente(valor: str) -> str:
    """Extrae solo el número de expediente en formato SIGLA NUMERO/AÑO.
    Ej: 'FMZ 041824/2019 (agregar/quitar...)' → 'FMZ 041824/2019'
         '041824/2019'                         → '041824/2019'
    """
    m = re.search(r"(?:[A-Za-z]{2,5}\s+)?\d+/\d{4}", valor)
    if m:
        return m.group(0).strip()
    return valor.strip()


def _docx_shade_cell(cell, fill: str) -> None:
    """Aplica color de fondo a una celda DOCX (fill sin #)."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def _docx_set_col_widths(table, widths_cm: list) -> None:
    """Establece el ancho de cada columna de una tabla DOCX."""
    for row in table.rows:
        for j, cell in enumerate(row.cells):
            if j < len(widths_cm):
                cell.width = DCm(widths_cm[j])


def _docx_add_table_header(table, headers: list, fill: str = "1E3A5F") -> None:
    """Escribe la fila de encabezado con fondo oscuro y texto blanco."""
    hdr_row = table.rows[0]
    for i, txt in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = ""
        _docx_shade_cell(cell, fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(txt)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(8)


def _docx_add_data_row(table, row_idx: int, values: list,
                       bold: bool = False, fill: str | None = None) -> None:
    """Escribe una fila de datos en la tabla DOCX."""
    row = table.rows[row_idx]
    for j, val in enumerate(values):
        cell = row.cells[j]
        cell.text = ""
        if fill:
            _docx_shade_cell(cell, fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(val))
        run.bold = bold
        run.font.size = Pt(8)


def generar_docx_cobro(
    resultado: dict,
    letrado: dict,
    caratula: str,
    expediente: str,
) -> bytes:
    """Genera escrito judicial DOCX — Intereses hasta el Cobro."""

    # ── Datos ────────────────────────────────────────────────────────────────
    caratula     = _limpiar_caratula(caratula)
    expediente   = limpiar_expediente(expediente)
    capital      = resultado["capital"]
    interes      = resultado["interes"]
    fecha_desde  = resultado["fecha_desde"]
    fecha_hasta  = resultado["fecha_hasta"]
    ind_ini      = resultado["indice_inicial"]
    ind_fin      = resultado["indice_final"]
    coef         = resultado["coeficiente"]

    monto_total    = interes
    nombre_letrado = letrado.get("nombre_completo", "").upper()
    cuil_letrado   = letrado.get("cuil", "")

    # ── Documento ────────────────────────────────────────────────────────────
    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = DCm(21.0)
    sec.page_height   = DCm(29.7)
    sec.left_margin   = DCm(3.0)
    sec.right_margin  = DCm(2.0)
    sec.top_margin    = DCm(2.5)
    sec.bottom_margin = DCm(2.5)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(12)

    def _ls(p) -> None:
        """Aplica interlineado 1.5 al párrafo."""
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    def _par_mixed(parts: list, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                   indent: float = 1.25, space_after: float = 6) -> None:
        """Párrafo con runs de distintos formatos. parts: [(text, bold, underline)]."""
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.first_line_indent = DCm(indent)
        p.paragraph_format.space_after = Pt(space_after)
        _ls(p)
        for text, bold, underline in parts:
            r = p.add_run(text)
            r.bold = bold
            r.underline = underline
            r.font.size = Pt(12)

    def _par(text: str, bold: bool = False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             indent: float = 1.25, space_after: float = 6) -> None:
        _par_mixed([(text, bold, False)], align=align, indent=indent, space_after=space_after)

    def _spacer(pt: float = 4) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(pt)
        p.paragraph_format.space_before = Pt(0)

    # ── TÍTULO ───────────────────────────────────────────────────────────────
    p_tit = doc.add_paragraph()
    p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_tit.paragraph_format.first_line_indent = DCm(0)
    p_tit.paragraph_format.space_after = Pt(8)
    _ls(p_tit)
    r = p_tit.add_run("PRACTICA LIQUIDACION – INTERESES")
    r.bold = True
    r.underline = True
    r.font.size = Pt(12)

    # ── SEÑOR JUEZ FEDERAL ───────────────────────────────────────────────────
    p_juez = doc.add_paragraph()
    p_juez.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_juez.paragraph_format.first_line_indent = DCm(0)
    p_juez.paragraph_format.space_after = Pt(6)
    _ls(p_juez)
    r2 = p_juez.add_run("SEÑOR JUEZ FEDERAL:")
    r2.bold = True
    r2.font.size = Pt(12)

    # ── Párrafo letrado + carátula ────────────────────────────────────────────
    # Carátula entre comillas dobles (no ángulos)
    p_intro = doc.add_paragraph()
    p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_intro.paragraph_format.first_line_indent = DCm(1.25)
    p_intro.paragraph_format.space_after = Pt(10)
    _ls(p_intro)
    for text, bold in [
        (nombre_letrado, True),
        (f', CUIL {cuil_letrado}, abogado, con personería acreditada en autos caratulados: "', False),
        (f"{caratula} - EXPTE. {expediente}", True),
        ('", con domicilio legal y electrónico constituido, a V.S. respetuosamente digo:', False),
    ]:
        r3 = p_intro.add_run(text)
        r3.bold = bold
        r3.font.size = Pt(12)

    # ── I. OBJETO ────────────────────────────────────────────────────────────
    _par_mixed([("I. OBJETO:", True, True)], indent=1.25, space_after=4)
    _par(
        "Que vengo por el presente a practicar liquidación de los intereses devengados "
        "hasta el efectivo pago del crédito reconocido en autos, conforme las pautas "
        "oportunamente establecidas y el criterio emergente de la sentencia firme.",
        space_after=10,
    )

    # ── II. LIQUIDACION PRACTICADA ────────────────────────────────────────────
    _par_mixed([("II. LIQUIDACION PRACTICADA:", True, True)], indent=1.25, space_after=4)
    _par(
        "Para la confección de la presente liquidación se aplicó la Tasa Pasiva Promedio "
        "del Banco Central de la República Argentina desde el día siguiente al que quedó "
        "aprobada la liquidación precedente y hasta la fecha de efectivo pago, entendiendo "
        "por tal aquella en que los fondos fueron efectivamente acreditados en la cuenta "
        "bancaria del actor, todo ello conforme constancias de autos.",
        space_after=6,
    )

    monto_palabras = _numero_a_palabras(monto_total)
    monto_numero   = _fmt(monto_total)
    _par_mixed([
        ("En consecuencia, la diferencia resultante a favor de la parte actora en concepto "
         "de intereses devengados hasta su efectivo pago asciende a la suma de PESOS ", False, False),
        (f"{monto_palabras} ($ {monto_numero})", True, False),
        (".", False, False),
    ], space_after=10)

    # ── PLANILLA ─────────────────────────────────────────────────────────────
    p_pl = doc.add_paragraph()
    p_pl.paragraph_format.first_line_indent = DCm(0)
    p_pl.paragraph_format.space_after = Pt(3)
    _ls(p_pl)
    rpl = p_pl.add_run("PLANILLA DE LIQUIDACIÓN")
    rpl.bold = True
    rpl.font.size = Pt(9)

    tbl = doc.add_table(rows=2, cols=7)
    tbl.style = "Table Grid"
    _docx_set_col_widths(tbl, [2.5, 2.3, 2.3, 2.0, 2.0, 2.4, 2.5])
    _docx_add_table_header(
        tbl,
        ["Capital ($)", "Int. desde", "Int. hasta",
         "Índice T₀", "Índice Tₘ", "Coeficiente", "Interés moratorio ($)"],
    )
    _docx_add_data_row(tbl, 1, [
        f"$ {_fmt(capital)}",
        fecha_desde.strftime("%d/%m/%Y"),
        fecha_hasta.strftime("%d/%m/%Y"),
        f"{ind_ini:,.4f}",
        f"{ind_fin:,.4f}",
        f"{coef:.6f}",
        f"$ {_fmt(interes)}",
    ], bold=True, fill="D8F3DC")
    _spacer(10)

    # ── III. DERECHO ─────────────────────────────────────────────────────────
    _par_mixed([("III. DERECHO:", True, True)], indent=1.25, space_after=4)
    _par(
        "La presente liquidación encuentra sustento en lo dispuesto en autos y en el "
        "criterio de sentencia firme que reconoce la procedencia de los intereses devengados "
        "hasta la íntegra satisfacción del crédito reconocido judicialmente.",
        space_after=6,
    )
    _par(
        "Asimismo, corresponde la aplicación de la tasa pasiva promedio del Banco Central "
        "de la República Argentina para el período comprendido entre la aprobación de la "
        "liquidación precedente y la efectiva acreditación de los fondos, en tanto durante "
        "dicho lapso subsistió impaga la obligación a cargo de la demandada.",
        space_after=10,
    )

    # ── IV. PETITUM ──────────────────────────────────────────────────────────
    _par_mixed([
        ("IV. PETITUM:", True, True),
        (" Por lo expuesto a V.S., solicito:", False, False),
    ], indent=1.25, space_after=4)
    _par(
        "a) Tenga por practicada la liquidación de intereses devengados hasta su efectivo "
        "pago conforme planilla que se acompaña.",
        space_after=4,
    )
    _par("b) Oportunamente, apruebe la misma en cuanto por derecho corresponda.", space_after=4)
    _par("c) Intime a la demandada al pago de las sumas resultantes.", space_after=10)

    # ── Cierre (centrado) ─────────────────────────────────────────────────────
    _par("Proveer de conformidad,", bold=True, indent=0, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    _par("SERÁ JUSTICIA.", bold=True, indent=0, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Intereses hasta el Cobro — PDF (Escrito judicial) ────────────────────────

def generar_pdf_cobro(
    resultado: dict,
    letrado: dict,
    caratula: str,
    expediente: str,
) -> bytes:
    """Genera escrito judicial en PDF — Intereses hasta el Cobro (mismo contenido que DOCX)."""
    from xml.sax.saxutils import escape as _xe

    caratula    = _limpiar_caratula(caratula)
    expediente  = limpiar_expediente(expediente)
    capital     = resultado["capital"]
    interes     = resultado["interes"]
    fecha_desde = resultado["fecha_desde"]
    fecha_hasta = resultado["fecha_hasta"]
    ind_ini     = resultado["indice_inicial"]
    ind_fin     = resultado["indice_final"]
    coef        = resultado["coeficiente"]

    monto_total    = interes
    nombre_letrado = letrado.get("nombre_completo", "").upper()
    cuil_letrado   = letrado.get("cuil", "")
    monto_palabras = _numero_a_palabras(monto_total)
    monto_numero   = _fmt(monto_total)

    FONT   = _PDF_FONT
    FONT_B = _PDF_FONT_BOLD
    LS     = 18   # 1.5 × 12pt

    def _ps(name, align=4, indent=35, space_after=6) -> ParagraphStyle:
        return ParagraphStyle(
            name, fontName=FONT, fontSize=12,
            leading=LS, alignment=align,
            firstLineIndent=indent, spaceAfter=space_after,
        )

    st_n  = _ps("cn")                           # normal justificado con sangría
    st_l  = _ps("cl", align=0, indent=0)        # izquierda
    st_c  = _ps("cc", align=1, indent=0, space_after=4)  # centrado

    buf = io.BytesIO()
    doc_pdf = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=3*cm, rightMargin=2*cm,
        topMargin=2.5*cm, bottomMargin=2.5*cm,
    )
    story = []

    # Título
    story.append(Paragraph(
        "<b><u>PRACTICA LIQUIDACION – INTERESES</u></b>",
        _ps("ct", align=1, indent=0, space_after=12),
    ))

    # SEÑOR JUEZ FEDERAL
    story.append(Paragraph(
        "<b>SEÑOR JUEZ FEDERAL:</b>",
        _ps("cj", align=0, indent=0, space_after=10),
    ))

    # Letrado + carátula
    story.append(Paragraph(
        f'<b>{_xe(nombre_letrado)}</b>'
        f', CUIL {_xe(cuil_letrado)}, abogado, con personería acreditada en autos '
        f'caratulados: "<b>{_xe(caratula)} - EXPTE. {_xe(expediente)}</b>", '
        f'con domicilio legal y electrónico constituido, a V.S. respetuosamente digo:',
        _ps("ci", space_after=10),
    ))

    # I. OBJETO
    story.append(Paragraph("<b><u>I. OBJETO:</u></b>", _ps("co1", align=0, indent=35)))
    story.append(Paragraph(
        "Que vengo por el presente a practicar liquidación de los intereses devengados "
        "hasta el efectivo pago del crédito reconocido en autos, conforme las pautas "
        "oportunamente establecidas y el criterio emergente de la sentencia firme.",
        _ps("co2", space_after=10),
    ))

    # II. LIQUIDACION PRACTICADA
    story.append(Paragraph("<b><u>II. LIQUIDACION PRACTICADA:</u></b>",
                           _ps("cl2", align=0, indent=35)))
    story.append(Paragraph(
        "Para la confección de la presente liquidación se aplicó la Tasa Pasiva Promedio "
        "del Banco Central de la República Argentina desde el día siguiente al que quedó "
        "aprobada la liquidación precedente y hasta la fecha de efectivo pago, entendiendo "
        "por tal aquella en que los fondos fueron efectivamente acreditados en la cuenta "
        "bancaria del actor, todo ello conforme constancias de autos.",
        st_n,
    ))
    story.append(Paragraph(
        f"En consecuencia, la diferencia resultante a favor de la parte actora en concepto "
        f"de intereses devengados hasta su efectivo pago asciende a la suma de PESOS "
        f"<b>{_xe(monto_palabras)} ($ {_xe(monto_numero)})</b>.",
        _ps("cl3", space_after=10),
    ))

    # PLANILLA DE LIQUIDACIÓN
    story.append(Paragraph(
        "<b>PLANILLA DE LIQUIDACIÓN</b>",
        _ps("cpl", align=0, indent=0, space_after=3),
    ))
    _headers = ["Capital ($)", "Int. desde", "Int. hasta",
                "Índice T₀", "Índice Tₘ", "Coeficiente", "Interés moratorio ($)"]
    _data = [_headers, [
        f"$ {_fmt(capital)}",
        fecha_desde.strftime("%d/%m/%Y"),
        fecha_hasta.strftime("%d/%m/%Y"),
        f"{ind_ini:,.4f}",
        f"{ind_fin:,.4f}",
        f"{coef:.6f}",
        f"$ {_fmt(interes)}",
    ]]
    _col_w = [w * cm for w in [2.5, 2.3, 2.3, 2.0, 2.0, 2.4, 2.5]]
    _tbl = Table(_data, colWidths=_col_w)
    _tbl.setStyle(TableStyle([
        ("BACKGROUND",    (0, 0), (-1, 0),  colors.HexColor("#1E3A5F")),
        ("TEXTCOLOR",     (0, 0), (-1, 0),  colors.white),
        ("FONTNAME",      (0, 0), (-1, 0),  FONT_B),
        ("FONTSIZE",      (0, 0), (-1, -1), 9),
        ("ALIGN",         (0, 0), (-1, -1), "CENTER"),
        ("BACKGROUND",    (0, 1), (-1, 1),  colors.HexColor("#D8F3DC")),
        ("FONTNAME",      (0, 1), (-1, 1),  FONT_B),
        ("GRID",          (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
        ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING",    (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(_tbl)
    story.append(Spacer(1, 0.4 * cm))

    # III. DERECHO
    story.append(Paragraph("<b><u>III. DERECHO:</u></b>",
                           _ps("cd1", align=0, indent=35)))
    story.append(Paragraph(
        "La presente liquidación encuentra sustento en lo dispuesto en autos y en el "
        "criterio de sentencia firme que reconoce la procedencia de los intereses devengados "
        "hasta la íntegra satisfacción del crédito reconocido judicialmente.",
        st_n,
    ))
    story.append(Paragraph(
        "Asimismo, corresponde la aplicación de la tasa pasiva promedio del Banco Central "
        "de la República Argentina para el período comprendido entre la aprobación de la "
        "liquidación precedente y la efectiva acreditación de los fondos, en tanto durante "
        "dicho lapso subsistió impaga la obligación a cargo de la demandada.",
        _ps("cd2", space_after=10),
    ))

    # IV. PETITUM
    story.append(Paragraph(
        "<b><u>IV. PETITUM:</u></b> Por lo expuesto a V.S., solicito:",
        _ps("cp1", align=0, indent=35),
    ))
    story.append(Paragraph(
        "a) Tenga por practicada la liquidación de intereses devengados hasta su efectivo "
        "pago conforme planilla que se acompaña.",
        st_n,
    ))
    story.append(Paragraph(
        "b) Oportunamente, apruebe la misma en cuanto por derecho corresponda.", st_n,
    ))
    story.append(Paragraph(
        "c) Intime a la demandada al pago de las sumas resultantes.",
        _ps("cp4", space_after=12),
    ))

    # Cierre centrado
    story.append(Paragraph("<b>Proveer de conformidad,</b>", st_c))
    story.append(Paragraph("<b>SERÁ JUSTICIA.</b>",
                           _ps("cfj", align=1, indent=0, space_after=0)))

    doc_pdf.build(story)
    return buf.getvalue()


# ── Ejecución de Sentencia — DOCX (Escrito judicial) ──────────────────────────

def generar_docx_ejecucion(
    resultado: dict,
    letrado: dict,
    caratula: str,
    expediente: str,
) -> bytes:
    """Genera escrito judicial DOCX — Ejecución de Sentencia (PASO #1)."""

    # ── Datos ────────────────────────────────────────────────────────────────
    caratula        = _limpiar_caratula(caratula)
    expediente      = limpiar_expediente(expediente)
    dia_120         = resultado["dia_120"]
    dia_121         = resultado["dia_121"]
    filas_a         = resultado["filas_a"]
    capital_a_total = resultado["capital_a_total"]
    res_a           = resultado["resultado_a"]
    res_b           = resultado["resultado_b"]
    fecha_hasta     = resultado["fecha_hasta"]

    int_a_val   = float(res_a.get("interes", 0) or 0) if not res_a.get("error") else 0.0
    int_b_val   = float(res_b["interes"].sum()) if not res_b.empty and "interes" in res_b.columns else 0.0
    monto_total = int_a_val + int_b_val

    nombre_letrado = letrado.get("nombre_completo", "").upper()
    cuil_letrado   = letrado.get("cuil", "")

    # ── Documento ────────────────────────────────────────────────────────────
    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = DCm(21.0)
    sec.page_height   = DCm(29.7)
    sec.left_margin   = DCm(3.0)
    sec.right_margin  = DCm(2.0)
    sec.top_margin    = DCm(2.5)
    sec.bottom_margin = DCm(2.5)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(12)

    def _ls(p) -> None:
        """Aplica interlineado 1.5 al párrafo."""
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    def _par_mixed(parts: list, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                   indent: float = 1.25, space_after: float = 6) -> None:
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.first_line_indent = DCm(indent)
        p.paragraph_format.space_after = Pt(space_after)
        _ls(p)
        for text, bold, underline in parts:
            r = p.add_run(text)
            r.bold = bold
            r.underline = underline
            r.font.size = Pt(12)

    def _par(text: str, bold: bool = False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             indent: float = 1.25, space_after: float = 6) -> None:
        _par_mixed([(text, bold, False)], align=align, indent=indent, space_after=space_after)

    def _section(num: str, titulo: str, space_after: float = 4) -> None:
        """Encabezado de sección: 'I.- OBJETO' en negrita."""
        _par_mixed([(f"{num} {titulo}", True, False)],
                   align=WD_ALIGN_PARAGRAPH.LEFT, indent=1.25, space_after=space_after)

    def _subsection(letra: str, titulo: str) -> None:
        """Subsección: 'A) Períodos anteriores...' en negrita."""
        _par_mixed([(f"{letra} {titulo}", True, False)],
                   align=WD_ALIGN_PARAGRAPH.LEFT, indent=1.25, space_after=4)

    def _spacer(pt: float = 4) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(pt)
        p.paragraph_format.space_before = Pt(0)

    # ── TÍTULO ───────────────────────────────────────────────────────────────
    p_tit = doc.add_paragraph()
    p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_tit.paragraph_format.first_line_indent = DCm(0)
    p_tit.paragraph_format.space_after = Pt(8)
    _ls(p_tit)
    r = p_tit.add_run("PRACTICA LIQUIDACION DE INTERESES MORATORIOS")
    r.bold = True
    r.underline = True
    r.font.size = Pt(12)

    # ── SEÑOR JUEZ FEDERAL ───────────────────────────────────────────────────
    p_juez = doc.add_paragraph()
    p_juez.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_juez.paragraph_format.first_line_indent = DCm(0)
    p_juez.paragraph_format.space_after = Pt(6)
    _ls(p_juez)
    rj = p_juez.add_run("SEÑOR JUEZ FEDERAL:")
    rj.bold = True
    rj.font.size = Pt(12)

    # ── Párrafo letrado + carátula ───────────────────────────────────────────
    p_intro = doc.add_paragraph()
    p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_intro.paragraph_format.first_line_indent = DCm(1.25)
    p_intro.paragraph_format.space_after = Pt(10)
    _ls(p_intro)
    for text, bold in [
        (nombre_letrado, True),
        (f', CUIL {cuil_letrado}, abogado, con personería acreditada en autos caratulados: "', False),
        (f"{caratula} - EXPTE. {expediente}", True),
        ('", con domicilio legal y electrónico constituido, a V.S. respetuosamente digo:', False),
    ]:
        ri = p_intro.add_run(text)
        ri.bold = bold
        ri.font.size = Pt(12)

    # ── I.- OBJETO ───────────────────────────────────────────────────────────
    _section("I.-", "OBJETO")
    _par(
        "Que vengo en legal tiempo y forma a acompañar nueva planilla de liquidación de "
        "intereses moratorios, practicada conforme los lineamientos expresamente establecidos "
        "por V.S. en fallo «Rastrilla» y «VEGA», solicitando oportunamente su aprobación.",
        space_after=10,
    )

    # ── II.- CUMPLIMIENTO ────────────────────────────────────────────────────
    _section("II.-", 'CUMPLIMIENTO DE LOS LINEAMIENTOS ESTABLECIDOS EN FALLO "VEGA" y "RASTRILLA"')
    _par(
        "Que la nueva liquidación acompañada ha sido confeccionada siguiendo estrictamente "
        "la metodología ordenada por V.S. en el considerando VII del decisorio referido.",
        space_after=6,
    )
    _par("En tal sentido, se procedió a distinguir expresamente:", space_after=6)

    # A) Períodos anteriores al día 120
    _subsection("A)", "Períodos anteriores al vencimiento del plazo de 120 días")
    _par_mixed([
        ("Respecto de los períodos comprendidos con anterioridad al vencimiento del plazo legal "
         "de ciento veinte (120) días hábiles para el cumplimiento de la sentencia, se determinó "
         "el retroactivo correspondiente por las diferencias devengadas hasta el día 120, es decir "
         "hasta el ", False, False),
        (dia_120.strftime("%d/%m/%Y"), True, False),
        (".", False, False),
    ], space_after=6)
    _par(
        "Posteriormente, sobre el monto total resultante, se calcularon intereses moratorios "
        "aplicando la tasa pasiva promedio del BCRA desde el día 121 —momento de constitución "
        "automática en mora conforme art. 886 CCCN— y hasta la fecha de efectiva transferencia "
        "del embargo.",
        space_after=6,
    )

    # B) Períodos posteriores al día 120
    _subsection("B)", "Períodos posteriores al vencimiento del plazo de 120 días")
    _par_mixed([
        ("Asimismo, para los períodos posteriores al vencimiento del referido plazo, se "
         "individualizó cada diferencia mensual devengada y se calcularon los intereses moratorios "
         "correspondientes desde que cada suma fue debida (", False, False),
        (dia_121.strftime("%d/%m/%Y"), True, False),
        (") y hasta la fecha de transferencia del embargo (", False, False),
        (fecha_hasta.strftime("%d/%m/%Y"), True, False),
        (").", False, False),
    ], space_after=6)
    _par(
        "De este modo, la metodología aplicada recepta íntegramente los parámetros fijados "
        "por V.S., respetando la diferenciación temporal expresamente establecida en la "
        "resolución dictada.",
        space_after=10,
    )

    # ── III.- PROCEDENCIA ────────────────────────────────────────────────────
    _section("III.-", "PROCEDENCIA DE LOS INTERESES MORATORIOS")
    _par(
        "Cabe destacar que V.S. ya ha reconocido expresamente la procedencia de los intereses "
        "moratorios reclamados, dejando establecido que la ANSES incurrió en mora automática "
        "una vez vencido el plazo de ciento veinte (120) días hábiles previsto para el "
        "cumplimiento de la sentencia.",
        space_after=6,
    )
    _par("En efecto, la resolución dictada en autos sostuvo expresamente que:", space_after=4)
    _par(
        "«a partir del día 121 —momento en el que se produce el incumplimiento de la "
        "sentencia—, la demandada ANSES se constituyó en mora en forma automática (conf. "
        "art. 886 del C.C.C.N) y debe intereses moratorios (conf. art. 768 del C.C.C.N).»",
        space_after=6,
    )
    _par(
        "Asimismo, V.S. dejó establecido que los mismos deben calcularse hasta la fecha de "
        "efectivo pago, extremo que ha sido debidamente respetado en la liquidación acompañada.",
        space_after=10,
    )

    # ── IV.- PLANILLA ────────────────────────────────────────────────────────
    _section("IV.-", "PLANILLA – ACOMPAÑA")

    monto_palabras = _numero_a_palabras(monto_total).lower()
    monto_numero   = _fmt(monto_total)
    _par_mixed([
        ("Que se acompaña planilla detallada de cálculo de intereses moratorios confeccionada "
         "conforme las pautas indicadas por V.S., discriminando períodos, capitales, fechas de "
         "mora, tasa aplicada hasta el momento de su efectivo pago. En consecuencia, resulta a "
         "favor de la actora monto de ", False, False),
        (f"pesos {monto_palabras} ($ {monto_numero})", True, False),
        (".", False, False),
    ], space_after=10)

    # — Tramo A —
    p_la = doc.add_paragraph()
    p_la.paragraph_format.first_line_indent = DCm(0)
    p_la.paragraph_format.space_after = Pt(3)
    _ls(p_la)
    ra2 = p_la.add_run("PLANILLA — TRAMO A")
    ra2.bold = True
    ra2.font.size = Pt(9)

    n_rows_a = 1 + len(filas_a) + 1
    tbl_a = doc.add_table(rows=n_rows_a, cols=2)
    tbl_a.style = "Table Grid"
    _docx_set_col_widths(tbl_a, [5.5, 10.5])
    _docx_add_table_header(tbl_a, ["Período", "Capital proporcional ($)"])
    for i, fila in enumerate(filas_a, start=1):
        _docx_add_data_row(tbl_a, i, [fila["periodo"], f"$ {_fmt(fila['capital'])}"])
    _docx_add_data_row(
        tbl_a, n_rows_a - 1,
        ["TOTAL TRAMO A", f"$ {_fmt(capital_a_total)}"],
        bold=True, fill="FEF3CD",
    )
    _spacer(4)

    if not res_a.get("error"):
        tbl_calc = doc.add_table(rows=2, cols=6)
        tbl_calc.style = "Table Grid"
        _docx_set_col_widths(tbl_calc, [3.0, 2.5, 2.5, 2.5, 2.5, 3.0])
        _docx_add_table_header(
            tbl_calc,
            ["Capital ($)", "Int. desde", "Int. hasta", "Índice T₀", "Índice Tₘ", "Interés moratorio ($)"],
            fill="2D6A4F",
        )
        _docx_add_data_row(tbl_calc, 1, [
            f"$ {_fmt(capital_a_total)}",
            dia_121.strftime("%d/%m/%Y"),
            fecha_hasta.strftime("%d/%m/%Y"),
            f"{res_a['indice_inicial']:,.4f}",
            f"{res_a['indice_final']:,.4f}",
            f"$ {_fmt(res_a['interes'])}",
        ], bold=True, fill="D8F3DC")
        _spacer(6)

    # — Tramo B —
    if not res_b.empty:
        df_b_ok = res_b[res_b["error"].isna()].copy()
        if not df_b_ok.empty:
            p_lb = doc.add_paragraph()
            p_lb.paragraph_format.first_line_indent = DCm(0)
            p_lb.paragraph_format.space_after = Pt(3)
            _ls(p_lb)
            rb2 = p_lb.add_run("PLANILLA — TRAMO B")
            rb2.bold = True
            rb2.font.size = Pt(9)

            n_rows_b = 1 + len(df_b_ok) + 1
            tbl_b = doc.add_table(rows=n_rows_b, cols=6)
            tbl_b.style = "Table Grid"
            _docx_set_col_widths(tbl_b, [2.2, 2.8, 2.5, 2.5, 3.0, 3.0])
            _docx_add_table_header(
                tbl_b,
                ["Período", "Capital ($)", "Int. desde", "Fecha pago", "Interés ($)", "Total ($)"],
            )
            for i, (_, row) in enumerate(df_b_ok.iterrows(), start=1):
                _docx_add_data_row(tbl_b, i, [
                    row["periodo"],
                    f"$ {_fmt(row['capital'])}",
                    row["fecha_desde"].strftime("%d/%m/%Y"),
                    row["fecha_pago"].strftime("%d/%m/%Y"),
                    f"$ {_fmt(row['interes'])}",
                    f"$ {_fmt(row['total'])}",
                ])
            t_int_b = df_b_ok["interes"].sum()
            t_tot_b = df_b_ok["total"].sum()
            _docx_add_data_row(
                tbl_b, n_rows_b - 1,
                ["TOTAL TRAMO B", "", "", "", f"$ {_fmt(t_int_b)}", f"$ {_fmt(t_tot_b)}"],
                bold=True, fill="FEF3CD",
            )
            _spacer(6)

    # Resultado final
    tbl_res = doc.add_table(rows=2, cols=2)
    tbl_res.style = "Table Grid"
    _docx_set_col_widths(tbl_res, [12.0, 4.0])
    _docx_add_table_header(
        tbl_res,
        ["RESULTADO FINAL — Intereses moratorios Tramo A + Tramo B", f"$ {_fmt(monto_total)}"],
        fill="052E16",
    )
    _docx_add_data_row(
        tbl_res, 1,
        ["Calculado hasta — Efectivo pago conforme recibo que consta en autos",
         fecha_hasta.strftime("%d/%m/%Y")],
        bold=True, fill="F0FDF4",
    )
    _spacer(10)

    # ── V.- PETITORIO ────────────────────────────────────────────────────────
    _par_mixed([
        ("V.-", True, False),
        (" PETITORIO: Por todo lo expuesto, a V.S. solicito:", False, False),
    ], align=WD_ALIGN_PARAGRAPH.LEFT, indent=1.25, space_after=4)
    _par(
        "1. Tenga por acompañada la nueva planilla de liquidación de intereses moratorios practicada.",
        space_after=4,
    )
    _par(
        "2. Oportunamente, apruebe la liquidación presentada en todas sus partes.",
        space_after=10,
    )

    # ── Cierre (centrado) ─────────────────────────────────────────────────────
    _par("Proveer de conformidad,", indent=0, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    _par("SERÁ JUSTICIA.", bold=True, indent=0, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Ejecución de Sentencia — PDF (Escrito judicial) ───────────────────────────

def generar_pdf_ejecucion_escrito(
    resultado: dict,
    letrado: dict,
    caratula: str,
    expediente: str,
) -> bytes:
    """Genera escrito judicial en PDF — Ejecución de Sentencia (mismo contenido que DOCX)."""
    from xml.sax.saxutils import escape as _xe

    caratula        = _limpiar_caratula(caratula)
    expediente      = limpiar_expediente(expediente)
    dia_120         = resultado["dia_120"]
    dia_121         = resultado["dia_121"]
    filas_a         = resultado["filas_a"]
    capital_a_total = resultado["capital_a_total"]
    res_a           = resultado["resultado_a"]
    res_b           = resultado["resultado_b"]
    fecha_hasta     = resultado["fecha_hasta"]

    int_a_val   = float(res_a.get("interes", 0) or 0) if not res_a.get("error") else 0.0
    int_b_val   = float(res_b["interes"].sum()) if not res_b.empty and "interes" in res_b.columns else 0.0
    monto_total = int_a_val + int_b_val

    nombre_letrado = letrado.get("nombre_completo", "").upper()
    cuil_letrado   = letrado.get("cuil", "")
    monto_palabras = _numero_a_palabras(monto_total).lower()
    monto_numero   = _fmt(monto_total)

    FONT   = _PDF_FONT
    FONT_B = _PDF_FONT_BOLD
    LS     = 18  # 1.5 × 12pt

    def _ps(name, align=4, indent=35, space_after=6) -> ParagraphStyle:
        return ParagraphStyle(
            name, fontName=FONT, fontSize=12,
            leading=LS, alignment=align,
            firstLineIndent=indent, spaceAfter=space_after,
        )

    buf = io.BytesIO()
    doc_pdf = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=3 * cm, rightMargin=2 * cm,
        topMargin=2.5 * cm, bottomMargin=2.5 * cm,
    )
    story = []

    # Título
    story.append(Paragraph(
        "<b><u>PRACTICA LIQUIDACION DE INTERESES MORATORIOS</u></b>",
        _ps("et", align=1, indent=0, space_after=12),
    ))

    # SEÑOR JUEZ FEDERAL
    story.append(Paragraph(
        "<b>SEÑOR JUEZ FEDERAL:</b>",
        _ps("ej", align=0, indent=0, space_after=10),
    ))

    # Letrado + carátula
    story.append(Paragraph(
        f'<b>{_xe(nombre_letrado)}</b>'
        f', CUIL {_xe(cuil_letrado)}, abogado, con personería acreditada en autos '
        f'caratulados: "<b>{_xe(caratula)} - EXPTE. {_xe(expediente)}</b>", '
        f'con domicilio legal y electrónico constituido, a V.S. respetuosamente digo:',
        _ps("ei", space_after=10),
    ))

    # I.- OBJETO
    story.append(Paragraph(
        "<b>I.-</b> OBJETO",
        _ps("eo1", align=0, indent=35),
    ))
    story.append(Paragraph(
        "Que vengo en legal tiempo y forma a acompañar nueva planilla de liquidación de "
        "intereses moratorios, practicada conforme los lineamientos expresamente establecidos "
        "por V.S. en fallo «Rastrilla» y «VEGA», solicitando oportunamente su aprobación.",
        _ps("eo2", space_after=10),
    ))

    # II.- CUMPLIMIENTO
    story.append(Paragraph(
        '<b>II.-</b> CUMPLIMIENTO DE LOS LINEAMIENTOS ESTABLECIDOS EN FALLO "VEGA" y "RASTRILLA"',
        _ps("ec1", align=0, indent=35),
    ))
    story.append(Paragraph(
        "Que la nueva liquidación acompañada ha sido confeccionada siguiendo estrictamente "
        "la metodología ordenada por V.S. en el considerando VII del decisorio referido.",
        _ps("ec2"),
    ))
    story.append(Paragraph(
        "En tal sentido, se procedió a distinguir expresamente:",
        _ps("ec3", space_after=6),
    ))

    # A)
    story.append(Paragraph(
        "<b>A) Períodos anteriores al vencimiento del plazo de 120 días</b>",
        _ps("ea1", align=0, indent=35),
    ))
    story.append(Paragraph(
        "Respecto de los períodos comprendidos con anterioridad al vencimiento del "
        "plazo legal de ciento veinte (120) días hábiles para el cumplimiento de la sentencia, "
        "se determinó el retroactivo correspondiente por las diferencias devengadas hasta el "
        f"día 120, es decir hasta el <b>{dia_120.strftime('%d/%m/%Y')}</b>.",
        _ps("ea2"),
    ))
    story.append(Paragraph(
        "Posteriormente, sobre el monto total resultante, se calcularon intereses moratorios "
        "aplicando la tasa pasiva promedio del BCRA desde el día 121 —momento de constitución "
        "automática en mora conforme art. 886 CCCN— y hasta la fecha de efectiva transferencia "
        "del embargo.",
        _ps("ea3"),
    ))

    # B)
    story.append(Paragraph(
        "<b>B) Períodos posteriores al vencimiento del plazo de 120 días</b>",
        _ps("eb1", align=0, indent=35),
    ))
    story.append(Paragraph(
        "Asimismo, para los períodos posteriores al vencimiento del referido plazo, se "
        "individualizó cada diferencia mensual devengada y se calcularon los intereses moratorios "
        f"correspondientes desde que cada suma fue debida (<b>{dia_121.strftime('%d/%m/%Y')}</b>) "
        f"y hasta la fecha de transferencia del embargo (<b>{fecha_hasta.strftime('%d/%m/%Y')}</b>).",
        _ps("eb2"),
    ))
    story.append(Paragraph(
        "De este modo, la metodología aplicada recepta íntegramente los parámetros fijados "
        "por V.S., respetando la diferenciación temporal expresamente establecida en la "
        "resolución dictada.",
        _ps("eb3", space_after=10),
    ))

    # III.- PROCEDENCIA
    story.append(Paragraph(
        "<b>III.-</b> PROCEDENCIA DE LOS INTERESES MORATORIOS",
        _ps("ep1", align=0, indent=35),
    ))
    story.append(Paragraph(
        "Cabe destacar que V.S. ya ha reconocido expresamente la procedencia de los "
        "intereses moratorios reclamados, dejando establecido que la ANSES incurrió en mora "
        "automática una vez vencido el plazo de ciento veinte (120) días hábiles previsto "
        "para el cumplimiento de la sentencia.",
        _ps("ep2"),
    ))
    story.append(Paragraph(
        "En efecto, la resolución dictada en autos sostuvo expresamente que:",
        _ps("ep3", space_after=4),
    ))
    story.append(Paragraph(
        "«a partir del día 121 —momento en el que se produce el incumplimiento de la "
        "sentencia—, la demandada ANSES se constituyó en mora en forma automática (conf. "
        "art. 886 del C.C.C.N) y debe intereses moratorios (conf. art. 768 del C.C.C.N).»",
        _ps("ep4", indent=70, space_after=6),
    ))
    story.append(Paragraph(
        "Asimismo, V.S. dejó establecido que los mismos deben calcularse hasta la fecha de "
        "efectivo pago, extremo que ha sido debidamente respetado en la liquidación acompañada.",
        _ps("ep5", space_after=10),
    ))

    # IV.- PLANILLA
    story.append(Paragraph(
        "<b>IV.-</b> PLANILLA – ACOMPAÑA",
        _ps("epl1", align=0, indent=35),
    ))
    story.append(Paragraph(
        "Que se acompaña planilla detallada de cálculo de intereses moratorios confeccionada "
        "conforme las pautas indicadas por V.S., discriminando períodos, capitales, fechas de "
        "mora, tasa aplicada hasta el momento de su efectivo pago. En consecuencia, resulta a "
        f"favor de la actora monto de <b>pesos {_xe(monto_palabras)} ($ {_xe(monto_numero)})</b>.",
        _ps("epl2", space_after=8),
    ))

    # Tramo A — períodos
    story.append(Paragraph("<b>PLANILLA — TRAMO A</b>",
                           _ps("eta", align=0, indent=0, space_after=3)))
    if filas_a:
        rows_a = [["Período", "Capital proporcional ($)"]]
        for f in filas_a:
            rows_a.append([f["periodo"], f"$ {_fmt(f['capital'])}"])
        rows_a.append(["TOTAL TRAMO A", f"$ {_fmt(capital_a_total)}"])
        t_a1 = Table(rows_a, colWidths=[4 * cm, 5 * cm], repeatRows=1)
        t_a1.setStyle(TableStyle([
            ("BACKGROUND",     (0, 0),  (-1, 0),  colors.HexColor("#1E3A5F")),
            ("TEXTCOLOR",      (0, 0),  (-1, 0),  colors.white),
            ("FONTNAME",       (0, 0),  (-1, 0),  FONT_B),
            ("FONTSIZE",       (0, 0),  (-1, -1), 9),
            ("ALIGN",          (1, 0),  (1, -1),  "RIGHT"),
            ("ROWBACKGROUNDS", (0, 1),  (-1, -2), [colors.white, colors.HexColor("#eef2f7")]),
            ("BACKGROUND",     (0, -1), (-1, -1), colors.HexColor("#fef3cd")),
            ("FONTNAME",       (0, -1), (-1, -1), FONT_B),
            ("GRID",           (0, 0),  (-1, -1), 0.5, colors.HexColor("#cccccc")),
            ("VALIGN",         (0, 0),  (-1, -1), "MIDDLE"),
            ("TOPPADDING",     (0, 0),  (-1, -1), 3),
            ("BOTTOMPADDING",  (0, 0),  (-1, -1), 3),
        ]))
        story.append(t_a1)
        story.append(Spacer(1, 0.2 * cm))

    # Tramo A — cálculo
    if not res_a.get("error"):
        rows_calc = [
            ["Capital ($)", "Int. desde", "Int. hasta", "Índice T₀", "Índice Tₘ", "Interés moratorio ($)"],
            [
                f"$ {_fmt(capital_a_total)}",
                dia_121.strftime("%d/%m/%Y"),
                fecha_hasta.strftime("%d/%m/%Y"),
                f"{res_a['indice_inicial']:,.4f}",
                f"{res_a['indice_final']:,.4f}",
                f"$ {_fmt(res_a['interes'])}",
            ],
        ]
        t_a2 = Table(rows_calc, colWidths=[w * cm for w in [3.0, 2.5, 2.5, 2.5, 2.5, 3.0]])
        t_a2.setStyle(TableStyle([
            ("BACKGROUND",    (0, 0), (-1, 0),  colors.HexColor("#2D6A4F")),
            ("TEXTCOLOR",     (0, 0), (-1, 0),  colors.white),
            ("FONTNAME",      (0, 0), (-1, 0),  FONT_B),
            ("FONTSIZE",      (0, 0), (-1, -1), 9),
            ("ALIGN",         (0, 0), (-1, -1), "CENTER"),
            ("BACKGROUND",    (0, 1), (-1, 1),  colors.HexColor("#D8F3DC")),
            ("FONTNAME",      (0, 1), (-1, 1),  FONT_B),
            ("GRID",          (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
            ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING",    (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]))
        story.append(t_a2)
        story.append(Spacer(1, 0.3 * cm))

    # Tramo B
    if not res_b.empty:
        df_b_ok = res_b[res_b["error"].isna()].copy()
        if not df_b_ok.empty:
            story.append(Paragraph("<b>PLANILLA — TRAMO B</b>",
                                   _ps("etb", align=0, indent=0, space_after=3)))
            rows_b = [["Período", "Capital ($)", "Int. desde", "Fecha pago", "Interés ($)", "Total ($)"]]
            for _, row in df_b_ok.iterrows():
                rows_b.append([
                    row["periodo"],
                    f"$ {_fmt(row['capital'])}",
                    row["fecha_desde"].strftime("%d/%m/%Y"),
                    row["fecha_pago"].strftime("%d/%m/%Y"),
                    f"$ {_fmt(row['interes'])}",
                    f"$ {_fmt(row['total'])}",
                ])
            t_ib = df_b_ok["interes"].sum()
            t_tb = df_b_ok["total"].sum()
            rows_b.append(["TOTAL TRAMO B", "", "", "",
                            f"$ {_fmt(t_ib)}", f"$ {_fmt(t_tb)}"])
            t_b = Table(rows_b, colWidths=[w * cm for w in [2.2, 2.8, 2.5, 2.5, 3.0, 3.0]],
                        repeatRows=1)
            t_b.setStyle(TableStyle([
                ("BACKGROUND",     (0, 0),  (-1, 0),  colors.HexColor("#1E3A5F")),
                ("TEXTCOLOR",      (0, 0),  (-1, 0),  colors.white),
                ("FONTNAME",       (0, 0),  (-1, 0),  FONT_B),
                ("FONTSIZE",       (0, 0),  (-1, -1), 9),
                ("ALIGN",          (0, 0),  (-1, -1), "CENTER"),
                ("ROWBACKGROUNDS", (0, 1),  (-1, -2), [colors.white, colors.HexColor("#eef2f7")]),
                ("BACKGROUND",     (0, -1), (-1, -1), colors.HexColor("#fef3cd")),
                ("FONTNAME",       (0, -1), (-1, -1), FONT_B),
                ("GRID",           (0, 0),  (-1, -1), 0.5, colors.HexColor("#cccccc")),
                ("VALIGN",         (0, 0),  (-1, -1), "MIDDLE"),
                ("TOPPADDING",     (0, 0),  (-1, -1), 3),
                ("BOTTOMPADDING",  (0, 0),  (-1, -1), 3),
            ]))
            story.append(t_b)
            story.append(Spacer(1, 0.3 * cm))

    # Resultado final
    _aw = (21.0 - 3.0 - 2.0) * cm
    rows_final = [
        ["RESULTADO FINAL — Intereses moratorios Tramo A + Tramo B", f"$ {_fmt(monto_total)}"],
        ["Calculado hasta — Efectivo pago conforme recibo que consta en autos",
         fecha_hasta.strftime("%d/%m/%Y")],
    ]
    t_fin = Table(rows_final, colWidths=[_aw * 0.72, _aw * 0.28])
    t_fin.setStyle(TableStyle([
        ("BACKGROUND",    (0, 0), (-1, 0),  colors.HexColor("#052E16")),
        ("TEXTCOLOR",     (0, 0), (-1, 0),  colors.white),
        ("FONTNAME",      (0, 0), (-1, 0),  FONT_B),
        ("FONTSIZE",      (0, 0), (-1, 0),  10),
        ("BACKGROUND",    (0, 1), (-1, 1),  colors.HexColor("#F0FDF4")),
        ("FONTNAME",      (0, 1), (-1, 1),  FONT_B),
        ("FONTSIZE",      (0, 1), (-1, 1),  9),
        ("ALIGN",         (1, 0), (1, -1),  "RIGHT"),
        ("GRID",          (0, 0), (-1, -1), 0.5, colors.HexColor("#16A34A")),
        ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
        ("TOPPADDING",    (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ("LEFTPADDING",   (0, 0), (-1, -1), 8),
    ]))
    story.append(t_fin)
    story.append(Spacer(1, 0.5 * cm))

    # V.- PETITORIO
    story.append(Paragraph(
        "<b>V.-</b> PETITORIO: Por todo lo expuesto, a V.S. solicito:",
        _ps("ev1", align=0, indent=35),
    ))
    story.append(Paragraph(
        "1. Tenga por acompañada la nueva planilla de liquidación de intereses moratorios practicada.",
        _ps("ev2"),
    ))
    story.append(Paragraph(
        "2. Oportunamente, apruebe la liquidación presentada en todas sus partes.",
        _ps("ev3", space_after=12),
    ))

    # Cierre centrado
    story.append(Paragraph("<b>Proveer de conformidad,</b>",
                           _ps("ecc", align=1, indent=0, space_after=4)))
    story.append(Paragraph("<b>SERÁ JUSTICIA.</b>",
                           _ps("ecj", align=1, indent=0, space_after=0)))

    doc_pdf.build(story)
    return buf.getvalue()


# ── Ampliación de Ejecución — DOCX (Escrito judicial) ─────────────────────────

def generar_docx_ampliacion(
    df_ok: "pd.DataFrame",
    letrado: dict,
    caratula: str,
    expediente: str,
) -> bytes:
    """Genera escrito judicial DOCX — Ampliación de Ejecución (PASO #2)."""

    # ── Datos ────────────────────────────────────────────────────────────────
    caratula   = _limpiar_caratula(caratula)
    expediente = limpiar_expediente(expediente)

    nombre_letrado = letrado.get("nombre_completo", "").upper()
    cuil_letrado   = letrado.get("cuil", "")

    t_cap = float(df_ok["capital"].sum())
    t_int = float(df_ok["interes"].sum())
    t_tot = float(df_ok["total"].sum())

    monto_palabras = _numero_a_palabras(t_int).lower()
    monto_numero   = _fmt(t_int)

    fecha_pago = df_ok["fecha_pago"].iloc[0] if "fecha_pago" in df_ok.columns else None
    fecha_pago_str = (
        fecha_pago.strftime("%d/%m/%Y")
        if fecha_pago is not None and hasattr(fecha_pago, "strftime")
        else str(fecha_pago) if fecha_pago is not None else ""
    )

    # Comillas tipográficas (curly quotes) — igual que el template Word
    QO = "“"  # "
    QC = "”"  # "

    # ── Documento ────────────────────────────────────────────────────────────
    doc = Document()
    sec = doc.sections[0]
    sec.page_width    = DCm(21.0)
    sec.page_height   = DCm(29.7)
    sec.left_margin   = DCm(3.0)
    sec.right_margin  = DCm(3.0)   # simétrico: igual que el template
    sec.top_margin    = DCm(2.5)
    sec.bottom_margin = DCm(2.5)
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(12)

    def _ls(p) -> None:
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    def _par_mixed(parts: list, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                   indent: float = 2.5, space_after: float = 6) -> None:
        """Párrafo con runs de distinto formato. indent = sangría primera línea en cm."""
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.first_line_indent = DCm(indent)
        p.paragraph_format.space_after = Pt(space_after)
        _ls(p)
        for text, bold, underline in parts:
            r = p.add_run(text)
            r.bold = bold
            r.underline = underline
            r.font.size = Pt(12)

    def _par(text: str, bold: bool = False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             indent: float = 2.5, space_after: float = 6) -> None:
        _par_mixed([(text, bold, False)], align=align, indent=indent, space_after=space_after)

    def _section(num: str, titulo: str, space_after: float = 4) -> None:
        """Título de sección: justificado, sangría 2.5 cm, negrita."""
        _par_mixed([(f"{num} {titulo}", True, False)],
                   align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=2.5, space_after=space_after)

    def _subsection(letra: str, titulo: str) -> None:
        """Subtítulo A)/B): justificado, sangría 2.5 cm, negrita."""
        _par_mixed([(f"{letra} {titulo}", True, False)],
                   align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=2.5, space_after=4)

    def _spacer(pt: float = 4) -> None:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(pt)
        p.paragraph_format.space_before = Pt(0)

    # ── TÍTULO ───────────────────────────────────────────────────────────────
    p_tit = doc.add_paragraph()
    p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_tit.paragraph_format.first_line_indent = DCm(0)
    p_tit.paragraph_format.space_after = Pt(8)
    _ls(p_tit)
    r = p_tit.add_run("PRACTICA AMPLIACION DE INTERESES MORATORIOS")
    r.bold = True
    r.underline = True
    r.font.size = Pt(12)

    # ── SEÑOR JUEZ FEDERAL ───────────────────────────────────────────────────
    p_juez = doc.add_paragraph()
    p_juez.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_juez.paragraph_format.first_line_indent = DCm(0)
    p_juez.paragraph_format.space_after = Pt(6)
    _ls(p_juez)
    rj = p_juez.add_run("SEÑOR JUEZ FEDERAL:")
    rj.bold = True
    rj.font.size = Pt(12)

    # ── Letrado + carátula ───────────────────────────────────────────────────
    p_intro = doc.add_paragraph()
    p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_intro.paragraph_format.first_line_indent = DCm(2.5)
    p_intro.paragraph_format.space_after = Pt(10)
    _ls(p_intro)
    for text, bold in [
        (nombre_letrado, True),
        (f", CUIL {cuil_letrado}, abogado, con personería acreditada en autos caratulados: {QO}", False),
        (f"{caratula} - EXPTE. {expediente}", True),
        (f"{QC}, con domicilio legal y electrónico constituido, a V.S. respetuosamente digo:", False),
    ]:
        ri = p_intro.add_run(text)
        ri.bold = bold
        ri.font.size = Pt(12)

    # ── I.- OBJETO ───────────────────────────────────────────────────────────
    _section("I.-", "OBJETO")
    _par(
        f"Que vengo en legal tiempo y forma a acompañar nueva planilla de liquidación de "
        f"intereses moratorios, practicada conforme los lineamientos expresamente establecidos "
        f"por V.S. en Fallo {QO}VEGA{QC}, solicitando oportunamente su aprobación.",
        space_after=10,
    )

    # ── II.- CUMPLIMIENTO ────────────────────────────────────────────────────
    _section("II.-",
             f"CUMPLIMIENTO DE LOS LINEAMIENTOS ESTABLECIDOS EN {QO}FALLO VEGA{QC} "
             f"y {QO}RASTRILLA{QC}")
    _par(
        "Que la nueva liquidación acompañada ha sido confeccionada siguiendo estrictamente "
        "la metodología ordenada por V.S. en el considerando VII del decisorio referido.",
        space_after=6,
    )
    _par("En tal sentido, se procedió a distinguir expresamente:", space_after=6)

    _subsection("A)", "Períodos anteriores al vencimiento del plazo de 120 días")
    _par(
        "Respecto de los períodos comprendidos con anterioridad al vencimiento del plazo legal "
        "de ciento veinte (120) días hábiles para el cumplimiento de la sentencia, se determinó "
        "el retroactivo correspondiente por las diferencias devengadas hasta el día 120.",
        space_after=6,
    )
    _par(
        "Posteriormente, sobre el monto total resultante, se calcularon intereses moratorios "
        "aplicando la tasa pasiva promedio del BCRA desde el día 121 —momento de constitución "
        "automática en mora conforme art. 886 CCCN— y hasta la fecha de efectiva transferencia "
        "del embargo.",
        space_after=6,
    )

    _subsection("B)", "Períodos posteriores al vencimiento del plazo de 120 días")
    _par(
        "Asimismo, para los períodos posteriores al vencimiento del referido plazo, se "
        "individualizó cada diferencia mensual devengada y se calcularon los intereses moratorios "
        "correspondientes desde que cada suma fue debida y hasta la fecha de transferencia del "
        "embargo.",
        space_after=6,
    )
    _par(
        "De este modo, la metodología aplicada recepta íntegramente los parámetros fijados "
        "por V.S., respetando la diferenciación temporal expresamente establecida en la "
        "resolución dictada.",
        space_after=10,
    )

    # ── III.- PROCEDENCIA ────────────────────────────────────────────────────
    _section("III.-", "PROCEDENCIA DE LOS INTERESES MORATORIOS")
    _par(
        "Cabe destacar que V.S. ya ha reconocido expresamente la procedencia de los intereses "
        "moratorios reclamados, dejando establecido que la ANSES incurrió en mora automática "
        "una vez vencido el plazo de ciento veinte (120) días hábiles previsto para el "
        "cumplimiento de la sentencia.",
        space_after=6,
    )
    _par("En efecto, la resolución dictada en autos sostuvo expresamente que:", space_after=4)

    # Cita textual: apertura normal + parte central en negrita+subrayado + cierre normal
    _par_mixed([
        (f"{QO}Por los periodos posteriores al vencimiento del plazo establecido en la "
         "sentencia de fondo para su cumplimiento (a partir del día 121), ", False, False),
        ("se deberán determinar las diferencias surgidas en cada mensual y proceder al cálculo "
         "de los intereses moratorios correspondientes desde que cada uno fue debido hasta la "
         "fecha de transferencia del embargo", True, True),
        (f"{QC}", False, False),
    ], space_after=6)

    _par_mixed([
        ("Asimismo, V.S. dejó establecido que los mismos deben calcularse hasta la fecha de "
         "efectivo pago ", False, False),
        (f"({fecha_pago_str})", True, False),
        (", extremo que ha sido debidamente respetado en la liquidación acompañada.", False, False),
    ], space_after=10)

    # ── IV.- PLANILLA ────────────────────────────────────────────────────────
    _section("IV.-", "PLANILLA – ACOMPAÑA")
    _par_mixed([
        ("Que se acompaña planilla de ampliación de intereses moratorios confeccionada "
         "conforme las pautas indicadas por V.S., discriminando períodos, capitales, fechas de "
         "mora, tasa aplicada y monto resultante, a saber de ", False, False),
        (f"pesos {monto_palabras} ($ {monto_numero})", True, False),
        (".", False, False),
    ], space_after=8)

    n_rows = 1 + len(df_ok) + 1
    tbl = doc.add_table(rows=n_rows, cols=8)
    tbl.style = "Table Grid"
    _docx_set_col_widths(tbl, [1.8, 2.2, 1.8, 1.9, 1.9, 2.1, 2.2, 2.1])
    _docx_add_table_header(tbl, [
        "Período", "Capital ($)", "Int. desde",
        "Índ. inicial", "Índ. final", "Coeficiente",
        "Interés ($)", "Total ($)",
    ])
    for i, (_, row) in enumerate(df_ok.iterrows(), start=1):
        _docx_add_data_row(tbl, i, [
            row["periodo"],
            f"$ {_fmt(row['capital'])}",
            row["fecha_desde"].strftime("%d/%m/%Y"),
            f"{row['indice_inicial']:,.4f}",
            f"{row['indice_final']:,.4f}",
            f"{row['coeficiente']:.6f}",
            f"$ {_fmt(row['interes'])}",
            f"$ {_fmt(row['total'])}",
        ])
    _docx_add_data_row(tbl, n_rows - 1, [
        "TOTAL", f"$ {_fmt(t_cap)}", "", "", "", "",
        f"$ {_fmt(t_int)}", f"$ {_fmt(t_tot)}",
    ], bold=True, fill="FEF3CD")
    _spacer(10)

    # ── V.- PETITORIO ────────────────────────────────────────────────────────
    _par_mixed([
        ("V.- PETITORIO:", True, False),
        (" Por todo lo expuesto, a V.S. solicito:", False, False),
    ], align=WD_ALIGN_PARAGRAPH.JUSTIFY, indent=2.5, space_after=4)
    _par("1. Tenga por acompañada la nueva planilla de liquidación de intereses moratorios "
         "practicada.", space_after=4)
    _par("2. Oportunamente, apruebe la liquidación presentada en todas sus partes.",
         space_after=10)

    # ── Cierre ───────────────────────────────────────────────────────────────
    _par("Proveer de conformidad,", bold=True, indent=0,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    _par("SERÁ JUSTICIA.", bold=True, indent=0,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Ampliación de Ejecución — PDF (Escrito judicial) ──────────────────────────

def generar_pdf_ampliacion(
    df_ok: "pd.DataFrame",
    letrado: dict,
    caratula: str,
    expediente: str,
) -> bytes:
    """Genera escrito judicial en PDF — Ampliación de Ejecución (PASO #2)."""
    from xml.sax.saxutils import escape as _xe

    caratula   = _limpiar_caratula(caratula)
    expediente = limpiar_expediente(expediente)

    nombre_letrado = letrado.get("nombre_completo", "").upper()
    cuil_letrado   = letrado.get("cuil", "")

    t_cap = float(df_ok["capital"].sum())
    t_int = float(df_ok["interes"].sum())
    t_tot = float(df_ok["total"].sum())

    monto_palabras = _numero_a_palabras(t_int).lower()
    monto_numero   = _fmt(t_int)

    fecha_pago = df_ok["fecha_pago"].iloc[0] if "fecha_pago" in df_ok.columns else None
    fecha_pago_str = (
        fecha_pago.strftime("%d/%m/%Y")
        if fecha_pago is not None and hasattr(fecha_pago, "strftime")
        else str(fecha_pago) if fecha_pago is not None else ""
    )

    FONT   = _PDF_FONT
    FONT_B = _PDF_FONT_BOLD
    LS     = 18  # 1.5 × 12pt

    def _ps(name, align=4, indent=35, space_after=6, left_indent=0, right_indent=0) -> ParagraphStyle:
        return ParagraphStyle(
            name, fontName=FONT, fontSize=12,
            leading=LS, alignment=align,
            firstLineIndent=indent, spaceAfter=space_after,
            leftIndent=left_indent, rightIndent=right_indent,
        )

    buf = io.BytesIO()
    doc_pdf = SimpleDocTemplate(
        buf, pagesize=A4,
        leftMargin=3 * cm, rightMargin=2 * cm,
        topMargin=2.5 * cm, bottomMargin=2.5 * cm,
    )
    story = []

    # Título
    story.append(Paragraph(
        "<b><u>PRACTICA AMPLIACION DE INTERESES MORATORIOS</u></b>",
        _ps("at", align=1, indent=0, space_after=12),
    ))

    # SEÑOR JUEZ FEDERAL
    story.append(Paragraph(
        "<b>SEÑOR JUEZ FEDERAL:</b>",
        _ps("aj", align=0, indent=0, space_after=10),
    ))

    # Letrado + carátula
    story.append(Paragraph(
        f'<b>{_xe(nombre_letrado)}</b>'
        f', CUIL {_xe(cuil_letrado)}, abogado, con personería acreditada en autos '
        f'caratulados: "<b>{_xe(caratula)} - EXPTE. {_xe(expediente)}</b>", '
        f'con domicilio legal y electrónico constituido, a V.S. respetuosamente digo:',
        _ps("ai", space_after=10),
    ))

    # I.- OBJETO
    story.append(Paragraph(
        '<b>I.-</b> OBJETO',
        _ps("ao1", align=1, indent=0),
    ))
    story.append(Paragraph(
        'Que vengo en legal tiempo y forma a acompañar nueva planilla de liquidación de '
        'intereses moratorios, practicada conforme los lineamientos expresamente establecidos '
        'por V.S. en Fallo "VEGA", solicitando oportunamente su aprobación.',
        _ps("ao2", space_after=10),
    ))

    # II.- CUMPLIMIENTO
    story.append(Paragraph(
        '<b>II.-</b> CUMPLIMIENTO DE LOS LINEAMIENTOS ESTABLECIDOS EN "FALLO VEGA" y "RASTRILLA"',
        _ps("ac1", align=1, indent=0),
    ))
    story.append(Paragraph(
        "Que la nueva liquidación acompañada ha sido confeccionada siguiendo estrictamente "
        "la metodología ordenada por V.S. en el considerando VII del decisorio referido.",
        _ps("ac2"),
    ))
    story.append(Paragraph(
        "En tal sentido, se procedió a distinguir expresamente:",
        _ps("ac3", space_after=6),
    ))
    story.append(Paragraph(
        "<b>A) Períodos anteriores al vencimiento del plazo de 120 días</b>",
        _ps("aa1", align=1, indent=0),
    ))
    story.append(Paragraph(
        "Respecto de los períodos comprendidos con anterioridad al vencimiento del plazo legal "
        "de ciento veinte (120) días hábiles para el cumplimiento de la sentencia, se determinó "
        "el retroactivo correspondiente por las diferencias devengadas hasta el día 120.",
        _ps("aa2"),
    ))
    story.append(Paragraph(
        "Posteriormente, sobre el monto total resultante, se calcularon intereses moratorios "
        "aplicando la tasa pasiva promedio del BCRA desde el día 121 —momento de constitución "
        "automática en mora conforme art. 886 CCCN— y hasta la fecha de efectiva transferencia "
        "del embargo.",
        _ps("aa3"),
    ))
    story.append(Paragraph(
        "<b>B) Períodos posteriores al vencimiento del plazo de 120 días</b>",
        _ps("ab1", align=1, indent=0),
    ))
    story.append(Paragraph(
        "Asimismo, para los períodos posteriores al vencimiento del referido plazo, se "
        "individualizó cada diferencia mensual devengada y se calcularon los intereses moratorios "
        "correspondientes desde que cada suma fue debida y hasta la fecha de transferencia del "
        "embargo.",
        _ps("ab2"),
    ))
    story.append(Paragraph(
        "De este modo, la metodología aplicada recepta íntegramente los parámetros fijados "
        "por V.S., respetando la diferenciación temporal expresamente establecida en la "
        "resolución dictada.",
        _ps("ab3", space_after=10),
    ))

    # III.- PROCEDENCIA
    story.append(Paragraph(
        "<b>III.-</b> PROCEDENCIA DE LOS INTERESES MORATORIOS",
        _ps("ap1", align=1, indent=0),
    ))
    story.append(Paragraph(
        "Cabe destacar que V.S. ya ha reconocido expresamente la procedencia de los intereses "
        "moratorios reclamados, dejando establecido que la ANSES incurrió en mora automática "
        "una vez vencido el plazo de ciento veinte (120) días hábiles previsto para el "
        "cumplimiento de la sentencia.",
        _ps("ap2"),
    ))
    story.append(Paragraph(
        "En efecto, la resolución dictada en autos sostuvo expresamente que:",
        _ps("ap3", space_after=4),
    ))
    story.append(Paragraph(
        '<b><u>"Por los periodos posteriores al vencimiento del plazo establecido en la sentencia '
        'de fondo para su cumplimiento (a partir del día 121), se deberán determinar las '
        'diferencias surgidas en cada mensual y proceder al cálculo de los intereses moratorios '
        'correspondientes desde que cada uno fue debido hasta la fecha de transferencia del embargo"</u></b>',
        _ps("ap4", indent=0, left_indent=57, right_indent=14, space_after=6),
    ))
    story.append(Paragraph(
        f"Asimismo, V.S. dejó establecido que los mismos deben calcularse hasta la fecha de "
        f"efectivo pago <b>({_xe(fecha_pago_str)})</b>, extremo que ha sido debidamente "
        f"respetado en la liquidación acompañada.",
        _ps("ap5", space_after=10),
    ))

    # IV.- PLANILLA
    story.append(Paragraph(
        "<b>IV.-</b> PLANILLA – ACOMPAÑA",
        _ps("apl1", align=1, indent=0),
    ))
    story.append(Paragraph(
        "Que se acompaña planilla de ampliación de intereses moratorios confeccionada "
        "conforme las pautas indicadas por V.S., discriminando períodos, capitales, fechas de "
        f"mora, tasa aplicada y monto resultante, a saber de "
        f"<b>pesos {_xe(monto_palabras)} ($ {_xe(monto_numero)})</b>.",
        _ps("apl2", space_after=8),
    ))

    _col_w = [w * cm for w in [1.8, 2.2, 1.8, 1.9, 1.9, 2.1, 2.2, 2.1]]
    _headers = ["Período", "Capital ($)", "Int. desde",
                "Índ. inicial", "Índ. final", "Coeficiente",
                "Interés ($)", "Total ($)"]
    rows_tbl = [_headers]
    for _, row in df_ok.iterrows():
        rows_tbl.append([
            row["periodo"],
            f"$ {_fmt(row['capital'])}",
            row["fecha_desde"].strftime("%d/%m/%Y"),
            f"{row['indice_inicial']:,.4f}",
            f"{row['indice_final']:,.4f}",
            f"{row['coeficiente']:.6f}",
            f"$ {_fmt(row['interes'])}",
            f"$ {_fmt(row['total'])}",
        ])
    rows_tbl.append([
        "TOTAL", f"$ {_fmt(t_cap)}", "", "", "", "",
        f"$ {_fmt(t_int)}", f"$ {_fmt(t_tot)}",
    ])
    t_planilla = Table(rows_tbl, colWidths=_col_w, repeatRows=1)
    t_planilla.setStyle(TableStyle([
        ("BACKGROUND",     (0, 0),  (-1, 0),  colors.HexColor("#1E3A5F")),
        ("TEXTCOLOR",      (0, 0),  (-1, 0),  colors.white),
        ("FONTNAME",       (0, 0),  (-1, 0),  FONT_B),
        ("FONTSIZE",       (0, 0),  (-1, -1), 8),
        ("ALIGN",          (0, 0),  (-1, -1), "CENTER"),
        ("ROWBACKGROUNDS", (0, 1),  (-1, -2), [colors.white, colors.HexColor("#eef2f7")]),
        ("BACKGROUND",     (0, -1), (-1, -1), colors.HexColor("#fef3cd")),
        ("FONTNAME",       (0, -1), (-1, -1), FONT_B),
        ("GRID",           (0, 0),  (-1, -1), 0.5, colors.HexColor("#cccccc")),
        ("VALIGN",         (0, 0),  (-1, -1), "MIDDLE"),
        ("TOPPADDING",     (0, 0),  (-1, -1), 3),
        ("BOTTOMPADDING",  (0, 0),  (-1, -1), 3),
    ]))
    story.append(t_planilla)
    story.append(Spacer(1, 0.5 * cm))

    # V.- PETITORIO
    story.append(Paragraph(
        "<b>V.- PETITORIO:</b> Por todo lo expuesto, a V.S. solicito:",
        _ps("av1", align=1, indent=0),
    ))
    story.append(Paragraph(
        "1. Tenga por acompañada la nueva planilla de liquidación de intereses moratorios practicada.",
        _ps("av2", indent=0, left_indent=35),
    ))
    story.append(Paragraph(
        "2. Oportunamente, apruebe la liquidación presentada en todas sus partes.",
        _ps("av3", indent=0, left_indent=35, space_after=12),
    ))

    # Cierre centrado
    story.append(Paragraph("<b>Proveer de conformidad,</b>",
                           _ps("acc", align=1, indent=0, space_after=4)))
    story.append(Paragraph("<b>SERÁ JUSTICIA.</b>",
                           _ps("acj", align=1, indent=0, space_after=0)))

    doc_pdf.build(story)
    return buf.getvalue()
