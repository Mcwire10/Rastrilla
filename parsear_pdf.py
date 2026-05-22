"""
Parsers específicos para cada formato de planilla de liquidación.
- parsear_bluecorp : PDF generado por BlueCorp
- parsear_jauregui : DOCX generado por Sistema Jauregui
- parsear_excel    : Excel o CSV (formato Jauregui o estándar)
"""
import io
import re
from datetime import date

import pandas as pd
import pdfplumber
from docx import Document
from docx.oxml.ns import qn

from calculos import primer_dia_mes_siguiente


# ── Helpers ───────────────────────────────────────────────────────────────────

def _monto(val) -> float | None:
    """Parsea un string de monto en formato AR (1.234,56 o 1234.56) → float."""
    if not val:
        return None
    val = re.sub(r'[^\d,.]', '', str(val)).strip()
    if not val:
        return None
    if ',' in val and '.' in val:
        val = val.replace('.', '').replace(',', '.')
    elif ',' in val:
        val = val.replace(',', '.')
    try:
        return float(val)
    except ValueError:
        return None


def _fecha(texto: str, patron: str) -> date | None:
    m = re.search(patron, texto, re.IGNORECASE)
    if not m:
        return None
    try:
        return pd.to_datetime(m.group(1), dayfirst=True, errors="coerce").date()
    except Exception:
        return None


def _fila(mes_str, anio_str, capital_str) -> dict | None:
    """Construye una fila a partir de mes, año y capital (string AR)."""
    if not str(mes_str).strip().isdigit():
        return None
    try:
        mes   = int(str(mes_str).strip())
        anio  = int(str(anio_str).strip())
        capital = _monto(capital_str)
    except (ValueError, TypeError):
        return None
    if capital is None or capital <= 0:
        return None
    periodo = f"{mes:02d}/{anio}"
    return {
        "periodo":     periodo,
        "capital":     capital,
        "fecha_desde": primer_dia_mes_siguiente(periodo),
        "fecha_pago":  None,
    }


def _fila_diff(mes_str, anio_str, percibido_str, reajustado_str) -> dict | None:
    """
    Construye una fila usando capital = Haber Reajustado − Haber Percibido.
    Descarta filas de encabezado, totales o valores no numéricos.
    """
    if not str(mes_str).strip().isdigit():
        return None
    try:
        mes        = int(str(mes_str).strip())
        anio       = int(str(anio_str).strip())
        percibido  = _monto(percibido_str)
        reajustado = _monto(reajustado_str)
    except (ValueError, TypeError):
        return None
    if percibido is None or reajustado is None:
        return None
    capital = round(reajustado - percibido, 2)
    if capital <= 0:
        return None
    periodo = f"{mes:02d}/{anio}"
    return {
        "periodo":     periodo,
        "capital":     capital,
        "fecha_desde": primer_dia_mes_siguiente(periodo),
        "fecha_pago":  None,
    }


# ── BlueCorp (PDF) ────────────────────────────────────────────────────────────

def parsear_bluecorp(file) -> pd.DataFrame:
    """
    PDF BlueCorp — sin límite de páginas.

    Estructura de columnas (pdfplumber, 12 cols):
      [0] Mes  [1] Año  [2] Haber Percibido  [3] Haber Reclamado/Reajustado
      [4] Diferencia Importe  [5] Dif.%  [6] Diferencia-Deducción
      [7] HAC  [8] Desc.OS  [9] Difer.+HAC-OS  [10] Capital  [11] Interés

    Capital = col[3] (Reajustado) − col[2] (Percibido)

    Página 1 : texto con fecha_pago ("calcularon hasta el DD/MM/YYYY")
    Páginas 2+: tablas de datos (soporta PDF de N páginas)
    """
    filas    = []
    vistos   = set()
    fecha_pago = None

    with pdfplumber.open(file) as pdf:
        # Fecha de pago en la primera página
        texto_p1 = pdf.pages[0].extract_text() or ""
        fecha_pago = _fecha(texto_p1, r'calcularon hasta el\s+(\d{1,2}/\d{1,2}/\d{4})')

        # Datos en páginas 2 en adelante (sin límite de páginas)
        for page in pdf.pages[1:]:
            for table in page.extract_tables():
                for row in table:
                    if not row or len(row) < 4:
                        continue
                    # Capital = Haber Reajustado (col[3]) − Haber Percibido (col[2])
                    f = _fila_diff(row[0], row[1], row[2], row[3])
                    if f and f["periodo"] not in vistos:
                        vistos.add(f["periodo"])
                        f["fecha_pago"] = fecha_pago
                        filas.append(f)

    if not filas:
        raise ValueError(
            "No se encontraron datos en el PDF BlueCorp.\n"
            "Verificá que el archivo sea una liquidación BlueCorp."
        )
    return pd.DataFrame(filas)


# ── Jauregui (DOCX) ──────────────────────────────────────────────────────────

def parsear_jauregui(file) -> pd.DataFrame:
    """
    DOCX Sistema Jauregui.

    Estructura de columnas en tabla anidada (15 cols):
      [0] Mes  [1] Año  [2] Percibido  [3] Reajustado  [4] Tope
      [5] %Conf  [6] Luego Tope  [7] Difer.  [8] SAC
      [9] Subtotal  [10] O.Social  [11] Dif.Neta  [12] Intereses
      [13] Total  [14] Texto

    Capital = col[3] (Reajustado) − col[2] (Percibido)

    Tabla principal fila 77 : fecha_pago
    Tabla anidada en fila 78: datos de períodos
    """
    filas      = []
    fecha_pago = None

    doc = Document(file)
    if len(doc.tables) < 2:
        raise ValueError("El DOCX Jauregui no tiene la estructura esperada.")

    main_table = doc.tables[1]

    # Fecha de pago en fila 77
    try:
        texto_77 = main_table.rows[77].cells[0].text
        fecha_pago = _fecha(texto_77, r'calcularon hasta el:\s*(\d{1,2}/\d{1,2}/\d{4})')
    except IndexError:
        pass

    # Tabla anidada en fila 78
    try:
        cell_78 = main_table.rows[78].cells[0]
    except IndexError:
        raise ValueError("No se encontró la sección de liquidación (fila 78) en el DOCX.")

    nested = cell_78._element.findall('.//' + qn('w:tbl'))
    if not nested:
        raise ValueError("No se encontró tabla de datos en el DOCX Jauregui.")

    vistos = set()
    for row_el in nested[0].findall('.//' + qn('w:tr')):
        cells = row_el.findall('.//' + qn('w:tc'))
        vals  = [''.join(t.text or '' for t in tc.findall('.//' + qn('w:t'))).strip()
                 for tc in cells]
        if len(vals) < 4:
            continue
        # Capital = Reajustado (col[3]) − Percibido (col[2])
        f = _fila_diff(vals[0], vals[1], vals[2], vals[3])
        if f:
            if f["periodo"] in vistos:
                continue
            vistos.add(f["periodo"])
            f["fecha_pago"] = fecha_pago
            filas.append(f)

    if not filas:
        raise ValueError(
            "No se encontraron datos en el DOCX Jauregui.\n"
            "Verificá que el archivo sea una liquidación del Sistema Jauregui."
        )
    return pd.DataFrame(filas)


# ── Excel / CSV ───────────────────────────────────────────────────────────────

def parsear_excel(file, csv: bool = False) -> pd.DataFrame:
    """
    Parsea Excel o CSV. Detecta dos formatos:

    1. Jauregui Excel: primera celda = "Mes", columnas "Percibido" y "Reajustado"
       Capital = Reajustado − Percibido  (columnas detectadas por nombre)
       Fallback: columna "Difer." o "Dif.Neta"

    2. Estándar: columnas periodo, capital, fecha_desde, fecha_pago
    """
    df_raw = pd.read_csv(file, header=None) if csv else pd.read_excel(file, header=None)

    if df_raw.empty:
        raise ValueError("El archivo está vacío.")

    first = df_raw.iloc[0].fillna("").astype(str).str.strip().str.lower().tolist()

    # Detectar formato Jauregui
    es_jauregui = first[0] == "mes" and any(
        v.startswith("difer") or ("dif" in v and "neta" in v)
        or "percibido" in v or "reajustado" in v
        for v in first if v
    )

    if es_jauregui:
        # Buscar columnas Percibido y Reajustado para calcular la diferencia directamente
        perc_idx  = next((i for i, v in enumerate(first) if "percibido"  in v), None)
        reaj_idx  = next((i for i, v in enumerate(first) if "reajustado" in v), None)

        # Si no existen, usar "Difer." o "Dif.Neta" como fallback
        if perc_idx is None or reaj_idx is None:
            try:
                dif_idx = next(
                    i for i, v in enumerate(first)
                    if v.startswith("difer") and "neta" not in v
                )
            except StopIteration:
                dif_idx = next(i for i, v in enumerate(first) if "dif" in v and "neta" in v)
            perc_idx = reaj_idx = None  # señal para usar dif_idx

        filas  = []
        vistos = set()
        for i in range(1, len(df_raw)):
            row     = df_raw.iloc[i]
            mes_str = str(row.iloc[0]).strip()
            if not mes_str.isdigit():
                continue
            try:
                anio_str = str(int(float(str(row.iloc[1]).strip())))
            except (ValueError, TypeError):
                continue

            if perc_idx is not None and reaj_idx is not None:
                f = _fila_diff(mes_str, anio_str,
                               str(row.iloc[perc_idx]), str(row.iloc[reaj_idx]))
            else:
                f = _fila(mes_str, anio_str, str(row.iloc[dif_idx]))

            if f is None or f["capital"] <= 0:
                continue
            if f["periodo"] in vistos:
                continue
            vistos.add(f["periodo"])
            filas.append(f)

        if not filas:
            raise ValueError("No se encontraron datos válidos en el Excel Jauregui.")
        return pd.DataFrame(filas)

    # ── Formato estándar ──────────────────────────────────────────────────────
    df = df_raw.copy()
    df.columns = df.iloc[0].fillna("").astype(str).str.lower().str.strip().str.replace(" ", "_")
    df = df.iloc[1:].reset_index(drop=True)
    mapa = {
        "período":  "periodo",  "periodo":  "periodo",
        "capital":  "capital",
        "dif._neta": "capital", "dif_neta": "capital", "dif.neta": "capital",
        "fecha_desde":     "fecha_desde", "desde":            "fecha_desde",
        "intereses_desde": "fecha_desde",
        "fecha_pago":      "fecha_pago",  "pago":             "fecha_pago",
    }
    df = df.rename(columns={c: mapa[c] for c in df.columns if c in mapa})
    for col in ["periodo", "capital", "fecha_desde", "fecha_pago"]:
        if col not in df.columns:
            df[col] = None
    df = df[["periodo", "capital", "fecha_desde", "fecha_pago"]]
    df["capital"]     = pd.to_numeric(df["capital"], errors="coerce")
    df["fecha_desde"] = pd.to_datetime(df["fecha_desde"], dayfirst=True, errors="coerce").dt.date
    df["fecha_pago"]  = pd.to_datetime(df["fecha_pago"],  dayfirst=True, errors="coerce").dt.date
    return df.dropna(subset=["periodo", "capital"])


# ── Jauregui (PDF) ───────────────────────────────────────────────────────────

def parsear_jauregui_pdf(file) -> pd.DataFrame:
    """
    PDF generado por el sistema Jauregui — sin límite de páginas.

    Estructura de columnas (16 cols, pdfplumber):
      [0] Mes  [1] Año  [2] Percibido  [3] Reajustado  [4] Tope
      [5] %Conf  [6] Luego Tope  [7] Difer.  [8] Dif.Act.  [9] SAC
      [10] Subtotal  [11] O.Social  [12] Dif.Neta  [13] Intereses
      [14] Total  [15] Texto

    Capital = col[3] (Reajustado) − col[2] (Percibido)

    Fecha hasta: texto "hasta el DD/MM/YYYY" o "calcularon hasta el: DD/MM/YYYY"
    La última fila es un total con col[2] vacío → se descarta automáticamente.
    """
    filas      = []
    vistos     = set()
    fecha_pago = None

    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            texto = page.extract_text() or ""
            if fecha_pago is None:
                fecha_pago = (
                    _fecha(texto, r'calcularon hasta el:\s*(\d{1,2}/\d{1,2}/\d{4})')
                    or _fecha(texto, r'hasta el\s+(\d{1,2}/\d{1,2}/\d{4})')
                )
            for table in page.extract_tables():
                for row in table:
                    if not row or len(row) < 4:
                        continue
                    # Capital = Reajustado (col[3]) − Percibido (col[2])
                    f = _fila_diff(row[0], row[1], row[2], row[3])
                    if f and f["periodo"] not in vistos:
                        vistos.add(f["periodo"])
                        f["fecha_pago"] = fecha_pago
                        filas.append(f)

    if not filas:
        raise ValueError(
            "No se encontraron datos en el PDF Jauregui.\n"
            "Verificá que el archivo sea una liquidación del Sistema Jauregui."
        )
    return pd.DataFrame(filas)


# ── Dispatcher ────────────────────────────────────────────────────────────────

def parsear_archivo(file, filename: str) -> tuple[pd.DataFrame, str]:
    """
    Detecta el formato por extensión (y para PDF también por contenido).
    Devuelve (df, nombre_formato).
    """
    ext = filename.lower().rsplit(".", 1)[-1]
    if ext == "pdf":
        # Leer contenido una sola vez para detectar el sistema y luego parsear
        content = file.read()
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            txt_p1 = pdf.pages[0].extract_text() or ""
        data = io.BytesIO(content)
        if "bluecorp" in txt_p1.lower():
            return parsear_bluecorp(data), "BlueCorp"
        else:
            return parsear_jauregui_pdf(data), "Jauregui PDF"
    if ext == "docx":
        return parsear_jauregui(file), "Jauregui DOCX"
    if ext in ("xlsx", "xls"):
        df = parsear_excel(file, csv=False)
        return df, "Jauregui Excel"
    if ext == "csv":
        df = parsear_excel(file, csv=True)
        return df, "CSV"
    raise ValueError(f"Formato no soportado: .{ext}")
